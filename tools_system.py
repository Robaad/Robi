"""
tools_system.py — Herramientas de sistema para Robi.

Cambios respecto a la versión anterior:
  - leer_eventos_calendar: ahora entiende 'mañana', 'ayer', nombres de días
    y fechas ISO. Logging detallado del rango resuelto.
  - crear_evento_calendar: validación más robusta de fecha/hora.
  - tavily_wait: sin cambios.
  - Google Calendar: sin cambios estructurales.
"""

from __future__ import annotations

import logging
import platform
import threading
import time
import re
from pathlib import Path
from datetime import datetime, timedelta
from zoneinfo import ZoneInfo

import requests
import whisper
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build
from google.auth.transport.requests import Request
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib.pyplot as plt


# ── Constantes ────────────────────────────────────────────────────────────────
TZ = ZoneInfo("Europe/Madrid")

SCOPES = ["https://www.googleapis.com/auth/calendar"]
service = None


# ── Tavily rate limiter (free tier: 1 req/seg) ────────────────────────────────
_TAVILY_MIN_INTERVAL = 1.2
_tavily_lock = threading.Lock()
_tavily_last_call: float = 0.0


def tavily_wait():
    """Throttle thread-safe para Tavily (free tier: 1 req/seg)."""
    global _tavily_last_call
    with _tavily_lock:
        ahora = time.monotonic()
        espera = _TAVILY_MIN_INTERVAL - (ahora - _tavily_last_call)
        if espera > 0:
            logging.debug("⏳ Tavily throttle: %.2fs", espera)
            time.sleep(espera)
        _tavily_last_call = time.monotonic()


# ── Whisper ───────────────────────────────────────────────────────────────────
if platform.system() == "Windows":
    WHISPER_CACHE = Path("./whisper_models")
else:
    WHISPER_CACHE = Path("/app/whisper_models")
WHISPER_CACHE.mkdir(exist_ok=True)


class LazyWhisperModel:
    """Carga Whisper bajo demanda (ahorra RAM en el arranque)."""

    def __init__(self, model_name: str = "base", cache_dir: Path = WHISPER_CACHE):
        self.model_name = model_name
        self.cache_dir = cache_dir
        self._model = None

    def _ensure(self):
        if self._model is None:
            logging.info("Cargando Whisper (%s)...", self.model_name)
            self._model = whisper.load_model(self.model_name, download_root=str(self.cache_dir))
        return self._model

    def transcribe(self, *args, **kwargs):
        return self._ensure().transcribe(*args, **kwargs)


modelo_whisper = LazyWhisperModel()


# ── Google Calendar ───────────────────────────────────────────────────────────

def init_calendar() -> bool:
    global service
    creds = None
    token_path = Path("token.json")

    if token_path.exists():
        creds = Credentials.from_authorized_user_file(str(token_path), SCOPES)

    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            try:
                creds.refresh(Request())
                logging.info("✅ Token Calendar refrescado")
            except Exception as e:
                logging.error("❌ Error refrescando token: %s", e)
                creds = None

        if not creds:
            if not Path("credentials.json").exists():
                logging.error("❌ Falta credentials.json")
                return False
            flow = InstalledAppFlow.from_client_secrets_file("credentials.json", SCOPES)
            creds = flow.run_local_server(port=0)

        token_path.write_text(creds.to_json())

    service = build("calendar", "v3", credentials=creds)
    logging.info("✅ Google Calendar conectado")
    return True


def _parsear_fecha_calendario(rango: str) -> tuple[datetime, datetime, str]:
    """
    Convierte 'rango' (string) en (inicio, fin, etiqueta) timezone-aware.

    Soporta:
      'hoy', 'today'
      'mañana', 'manana', 'tomorrow'
      'ayer', 'yesterday'
      nombres de días: 'lunes', 'martes', ..., 'domingo'
        → próximo día de esa semana (si hoy es ese día, el siguiente)
      'semana', 'esta semana', 'week' → próximos 7 días
      'próxima semana', 'proxima semana' → semana siguiente completa
      'YYYY-MM-DD' → ese día concreto
      cualquier otra cosa → hoy (con aviso en log)
    """
    ahora = datetime.now(TZ)
    hoy = ahora.replace(hour=0, minute=0, second=0, microsecond=0)
    r = rango.strip().lower()

    # ── Hoy ──────────────────────────────────────────────────────────────
    if r in ("hoy", "today", ""):
        inicio = hoy
        fin = hoy + timedelta(days=1) - timedelta(seconds=1)
        return inicio, fin, "hoy"

    # ── Mañana ───────────────────────────────────────────────────────────
    if r in ("mañana", "manana", "tomorrow"):
        inicio = hoy + timedelta(days=1)
        fin = inicio + timedelta(days=1) - timedelta(seconds=1)
        return inicio, fin, "mañana"

    # ── Ayer ─────────────────────────────────────────────────────────────
    if r in ("ayer", "yesterday"):
        inicio = hoy - timedelta(days=1)
        fin = hoy - timedelta(seconds=1)
        return inicio, fin, "ayer"

    # ── Nombres de días de la semana ──────────────────────────────────────
    DIAS = {
        "lunes": 0, "martes": 1, "miércoles": 2, "miercoles": 2,
        "jueves": 3, "viernes": 4, "sábado": 5, "sabado": 5, "domingo": 6,
        "monday": 0, "tuesday": 1, "wednesday": 2, "thursday": 3,
        "friday": 4, "saturday": 5, "sunday": 6,
    }
    if r in DIAS:
        objetivo = DIAS[r]
        dias_hasta = (objetivo - ahora.weekday()) % 7
        if dias_hasta == 0:
            dias_hasta = 7  # Si es hoy, buscar el de la semana que viene
        inicio = hoy + timedelta(days=dias_hasta)
        fin = inicio + timedelta(days=1) - timedelta(seconds=1)
        etiqueta = inicio.strftime("%A %d/%m/%Y")
        return inicio, fin, etiqueta

    # ── Semana ────────────────────────────────────────────────────────────
    if r in ("semana", "esta semana", "week", "próximos 7 días", "proximos 7 dias"):
        inicio = ahora
        fin = ahora + timedelta(days=7)
        return inicio, fin, "próximos 7 días"

    # ── Próxima semana ────────────────────────────────────────────────────
    if r in ("próxima semana", "proxima semana", "next week"):
        # Lunes de la semana que viene
        dias_hasta_lunes = (7 - ahora.weekday()) % 7 or 7
        inicio = hoy + timedelta(days=dias_hasta_lunes)
        fin = inicio + timedelta(days=7)
        return inicio, fin, f"semana del {inicio.strftime('%d/%m')}"

    # ── Fecha ISO YYYY-MM-DD ──────────────────────────────────────────────
    iso_match = re.match(r"^(\d{4})-(\d{2})-(\d{2})$", r)
    if iso_match:
        try:
            inicio = datetime(
                int(iso_match.group(1)),
                int(iso_match.group(2)),
                int(iso_match.group(3)),
                tzinfo=TZ,
            )
            fin = inicio + timedelta(days=1) - timedelta(seconds=1)
            return inicio, fin, inicio.strftime("%d/%m/%Y")
        except ValueError as e:
            logging.warning("Fecha ISO inválida '%s': %s", rango, e)

    # ── Fallback: hoy ─────────────────────────────────────────────────────
    logging.warning("Rango de calendario no reconocido '%s' → usando 'hoy'", rango)
    inicio = hoy
    fin = hoy + timedelta(days=1) - timedelta(seconds=1)
    return inicio, fin, f"hoy (rango '{rango}' no reconocido)"


def leer_eventos_calendar(rango: str = "semana") -> str:
    """
    Lee eventos de Google Calendar.
    rango: 'hoy', 'mañana', 'semana', nombre de día, o 'YYYY-MM-DD'.
    """
    if not service:
        return "❌ Google Calendar no está inicializado. Comprueba el token."

    try:
        inicio, fin, etiqueta = _parsear_fecha_calendario(rango)
        logging.info(
            "📅 Calendar: rango='%s' → %s .. %s",
            rango,
            inicio.isoformat(),
            fin.isoformat(),
        )

        result = (
            service.events()
            .list(
                calendarId="primary",
                timeMin=inicio.isoformat(),
                timeMax=fin.isoformat(),
                maxResults=20,
                singleEvents=True,
                orderBy="startTime",
            )
            .execute()
        )
        events = result.get("items", [])

        if not events:
            return f"📅 No hay eventos para {etiqueta}."

        lineas = [f"📅 Eventos — {etiqueta}:\n"]
        for ev in events:
            start_raw = ev["start"].get("dateTime", ev["start"].get("date", ""))
            try:
                dt = datetime.fromisoformat(
                    start_raw.replace("Z", "+00:00")
                ).astimezone(TZ)
                hora_str = dt.strftime("%d/%m %H:%M")
            except Exception:
                hora_str = start_raw
            lineas.append(f"• {hora_str} — {ev.get('summary', 'Sin título')}")

        return "\n".join(lineas)

    except Exception as e:
        logging.error("Error leyendo Calendar: %s", e)
        return f"❌ Error al leer el calendario: {e}"


def crear_evento_calendar(titulo: str, fecha: str, hora: str) -> str:
    """
    Crea un evento en Google Calendar.
    fecha: YYYY-MM-DD  |  hora: HH:MM
    """
    if not service:
        return "❌ Google Calendar no está inicializado."

    # Validar formato de fecha
    if not re.match(r"^\d{4}-\d{2}-\d{2}$", fecha):
        return (
            f"❌ Formato de fecha incorrecto: '{fecha}'. "
            "Usa YYYY-MM-DD (ej: 2026-04-15)."
        )
    # Validar formato de hora
    if not re.match(r"^\d{2}:\d{2}$", hora):
        return (
            f"❌ Formato de hora incorrecto: '{hora}'. "
            "Usa HH:MM en 24h (ej: 18:30)."
        )

    try:
        start_dt = f"{fecha}T{hora}:00"
        end_dt_obj = datetime.fromisoformat(f"{fecha}T{hora}:00").replace(tzinfo=TZ)
        end_dt = (end_dt_obj + timedelta(minutes=30)).strftime("%Y-%m-%dT%H:%M:00")

        event = {
            "summary": titulo,
            "start": {"dateTime": start_dt, "timeZone": "Europe/Madrid"},
            "end":   {"dateTime": end_dt,   "timeZone": "Europe/Madrid"},
            "reminders": {
                "useDefault": False,
                "overrides": [{"method": "popup", "minutes": 30}],
            },
        }
        service.events().insert(calendarId="primary", body=event).execute()
        logging.info("✅ Evento creado: '%s' el %s a las %s", titulo, fecha, hora)
        return f"✅ Evento '{titulo}' creado para el {fecha} a las {hora}h (aviso 30 min antes)."

    except Exception as e:
        logging.error("Error creando evento Calendar: %s", e)
        return f"❌ Error al crear el evento: {e}"


# ── Domótica OpenHAB ──────────────────────────────────────────────────────────
GRUPOS_DOMOTICA = {
    "banyo 0": ["Banyo1", "Espejo1"],
    "banyo 1": ["BanyoP1_Button1", "BanyoP1_Button2"],
    "banyo 2": ["banyoP2", "Espejo2"],
    "comedor": ["Comedor", "PasilloSur", "Lampara", "Siri"],
    "dormitorio": ["Habitacio", "Led", "Vestidor"],
    "salon": ["Biblioteca", "Pared", "Piano", "Salon_Button1"],
    "despacho": ["Despacho"],
    "cocina": ["Cocina"],
    "ovidi": ["Ovidi"],
    "caterina": ["Caterina"],
    "ventilador caterina": ["ventiladorCaterina"],
    "ventilador ovidi": ["ventiladorOvidi"],
    "ventilador dormitori": ["HabitacioVentilador"],
    "ventilador despacho": ["ventiladorDespacho"],
}


def control_openhab(item: str, state: str, config: dict) -> str:
    """Controla dispositivos OpenHAB. state: 'ON' o 'OFF'."""
    state_final = "ON" if "ON" in state.upper() else "OFF"
    item_key = item.lower().strip()
    items = GRUPOS_DOMOTICA.get(item_key, [item])

    errores = []
    for i in items:
        try:
            r = requests.post(
                f"{config['openhab']['url']}/items/{i}",
                data=state_final,
                headers={"Content-Type": "text/plain"},
                timeout=5,
            )
            if r.status_code not in (200, 202):
                errores.append(f"{i}: HTTP {r.status_code}")
            else:
                logging.info("✅ OpenHAB: %s → %s", i, state_final)
        except Exception as e:
            errores.append(f"{i}: {e}")
            logging.error("❌ OpenHAB %s: %s", i, e)

    if errores:
        return f"⚠️ {item} → {state_final} (errores: {', '.join(errores)})"
    return f"✅ {item} → {state_final}"


# ── IP Pública ────────────────────────────────────────────────────────────────
def obtener_ip_publica() -> str:
    servicios = [
        ("https://api.ipify.org?format=json", lambda r: r.json()["ip"]),
        ("https://api.myip.com",              lambda r: r.json()["ip"]),
        ("https://ifconfig.me/ip",            lambda r: r.text.strip()),
    ]
    for url, extractor in servicios:
        try:
            r = requests.get(url, timeout=5)
            r.raise_for_status()
            ip = extractor(r)
            logging.info("✅ IP pública: %s", ip)
            return f"🌐 Tu IP pública es: **{ip}**"
        except Exception:
            continue
    return "❌ No pude obtener la IP pública"


# ── Búsqueda Web ──────────────────────────────────────────────────────────────
def buscar_internet(query: str, client, config: dict, modelo: str) -> str:
    """Busca en Tavily y sintetiza el resultado con el LLM."""
    try:
        query = query.strip("'\" \n")
        if not query:
            return "❌ La consulta de búsqueda está vacía."

        tavily_wait()

        r = requests.post(
            "https://api.tavily.com/search",
            json={
                "api_key": config["tavily"]["api_key"],
                "query": query,
                "search_depth": "basic",
                "max_results": 3,
            },
            timeout=10,
        )
        r.raise_for_status()
        results = r.json().get("results", [])

        if not results:
            return "❌ Sin resultados en la web."

        contenido = "\n\n".join(
            f"Fuente {i + 1}: {x['content']}" for i, x in enumerate(results)
        )

        prompt = (
            f"Responde de forma concisa y directa a: \"{query}\"\n\n"
            f"Fuentes:\n{contenido}\n\n"
            "Responde en máximo 3 frases."
        )
        resp = client.chat.complete(
            model=modelo,
            messages=[{"role": "user", "content": prompt}],
        )
        return resp.choices[0].message.content or "Sin respuesta."

    except Exception as e:
        logging.error("Error en búsqueda Tavily: %s", e)
        return f"❌ Error de búsqueda: {e}"


# ── Exportar a Word ───────────────────────────────────────────────────────────
def exportar_a_word_premium(estudio_data: dict, nombre_archivo="Informe_Robi_Pro.docx") -> str:
    """Convierte la salida del ContentEngine en un Word profesional."""
    import io
    doc = Document()

    titulo_principal = estudio_data["metadata"]["tema"].upper()
    p = doc.add_heading(titulo_principal, 0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        f"\nFecha: {datetime.now().strftime('%d/%m/%Y')}\n"
        f"Nivel: {estudio_data['metadata']['nivel']}\n\n"
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    doc.add_heading("Índice de Contenidos", level=1)
    for titulo in estudio_data["indice"]:
        doc.add_paragraph(titulo, style="List Bullet")
    doc.add_page_break()

    for sec in estudio_data["secciones"]:
        doc.add_heading(f"{sec['numero']}. {sec['titulo']}", level=1)
        doc.add_paragraph(sec["contenido"])

        if sec.get("datos_visuales"):
            vis = sec["datos_visuales"]
            tipo = vis.get("tipo", "barras")
            try:
                if tipo == "organigrama":
                    doc.add_heading(f"Diagrama: {vis['titulo']}", level=3)
                    for padre, hijos in vis["datos"].items():
                        doc.add_paragraph(f"▪  {padre}", style="List Bullet")
                        for hijo in hijos:
                            doc.add_paragraph(hijo, style="List Continue Bullet")
                elif tipo in ("barras", "tarta", "lineas"):
                    plt.figure(figsize=(8, 5))
                    etiquetas = list(vis["datos"].keys())
                    valores = list(vis["datos"].values())
                    if tipo == "barras":
                        plt.bar(etiquetas, valores, color="#2E86C1")
                    elif tipo == "tarta":
                        plt.pie(
                            valores, labels=etiquetas, autopct="%1.1f%%",
                            colors=["#2E86C1", "#AED6F1", "#1B4F72"],
                        )
                    plt.title(vis["titulo"])
                    buf = io.BytesIO()
                    plt.savefig(buf, format="png", bbox_inches="tight")
                    plt.close()
                    doc.add_picture(buf, width=Inches(5.5))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                logging.error("Error en gráfico sección %s: %s", sec["numero"], e)

    doc.save(nombre_archivo)
    return nombre_archivo