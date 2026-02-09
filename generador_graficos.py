"""
GENERADOR DE GRÁFICOS - Extracción de datos y visualización automática
======================================================================
Este módulo analiza texto académico, extrae datos numéricos y genera
gráficos profesionales para insertar en documentos Word.
"""

import re
import logging
import asyncio
import matplotlib.pyplot as plt
import matplotlib
from matplotlib import patches
matplotlib.use('Agg')  # Backend sin GUI para servidores
from io import BytesIO
from typing import Dict, List, Optional, Tuple
import numpy as np

# Configuración de estilo profesional
plt.style.use('seaborn-v0_8-darkgrid')
COLORES_PROFESIONALES = ['#2E86C1', '#E74C3C', '#27AE60', '#F39C12', '#8E44AD', '#16A085']


class ExtractorDatos:
    """Extrae datos estructurados de texto académico."""
    
    def __init__(self, client, modelo="mistral-large-latest"):
        self.client = client
        self.modelo = modelo
    
    async def detectar_y_extraer_datos(self, contenido: str, titulo_seccion: str) -> Optional[Dict]:
        """
        Detecta si una sección contiene datos visualizables y los extrae.
        
        Returns:
            Dict con 'tipo', 'titulo', 'datos' si encuentra datos
            None si no hay datos visualizables
        """
        
        # Detectar si el contenido tiene datos numéricos relevantes
        indicadores_datos = [
            r'\d+%',  # Porcentajes
            r'\d+\.\d+',  # Decimales
            r'\d{1,3}(,\d{3})*',  # Números con separadores de miles
            'estadística', 'porcentaje', 'distribución', 'comparación',
            'tabla', 'datos', 'cifras', 'número', 'cantidad'
        ]
        
        tiene_datos = any(
            re.search(patron, contenido, re.IGNORECASE) 
            for patron in indicadores_datos
        )
        
        if not tiene_datos:
            return None
        
        # Usar IA para extraer datos estructurados
        prompt = f"""Analiza este texto de la sección "{titulo_seccion}" y extrae SOLO datos numéricos visualizables.

TEXTO:
{contenido[:1500]}

REGLAS:
1. Si encuentras datos numéricos comparables (porcentajes, cantidades, distribuciones), extráelos
2. Responde SOLO en formato JSON
3. Si NO hay datos visualizables, responde: {{"tiene_datos": false}}
4. Si hay datos, responde:
   {{
     "tiene_datos": true,
     "tipo_grafico": "barras|lineas|tarta|puntos",
     "titulo": "Título descriptivo del gráfico",
     "eje_x": "Nombre del eje X",
     "eje_y": "Nombre del eje Y",
     "datos": {{
       "Categoría 1": valor_numerico,
       "Categoría 2": valor_numerico,
       ...
     }}
   }}

EJEMPLOS:

Texto: "El 45% prefiere Python, 30% Java y 25% JavaScript"
Respuesta: {{"tiene_datos": true, "tipo_grafico": "tarta", "titulo": "Preferencias de lenguajes", "datos": {{"Python": 45, "Java": 30, "JavaScript": 25}}}}

Texto: "La adopción de IA creció del 20% en 2020 al 65% en 2024"
Respuesta: {{"tiene_datos": true, "tipo_grafico": "lineas", "titulo": "Crecimiento adopción IA", "eje_x": "Año", "eje_y": "Porcentaje", "datos": {{"2020": 20, "2021": 30, "2022": 42, "2023": 55, "2024": 65}}}}

Responde SOLO el JSON, sin explicaciones.
"""
        
        try:
            response = await asyncio.to_thread(
                self.client.chat.complete,
                model=self.modelo,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.1,
                response_format={"type": "json_object"}
            )
            
            import json
            resultado = json.loads(response.choices[0].message.content)
            
            if not resultado.get('tiene_datos', False):
                return None
            
            return {
                'tipo': resultado['tipo_grafico'],
                'titulo': resultado['titulo'],
                'eje_x': resultado.get('eje_x', ''),
                'eje_y': resultado.get('eje_y', 'Valor'),
                'datos': resultado['datos']
            }
        
        except Exception as e:
            logging.error(f"Error extrayendo datos: {e}")
            return None


class ExtractorVisuales:
    """Detecta oportunidades para ilustraciones y diagramas a partir de texto."""

    def __init__(self, client, modelo="mistral-large-latest"):
        self.client = client
        self.modelo = modelo

    async def detectar_visual(self, contenido: str, titulo_seccion: str) -> Optional[Dict]:
        """
        Detecta si un texto se presta a un diagrama o ilustración conceptual.

        Returns:
            Dict con especificación visual si aplica; None si no aplica.
        """
        indicadores_visual = [
            'proceso', 'flujo', 'etapas', 'fases', 'pipeline', 'metodología',
            'arquitectura', 'componentes', 'modelo', 'framework', 'relación',
            'mapa', 'conceptual', 'ciclo', 'sistema'
        ]

        tiene_indicios = any(
            palabra in contenido.lower() for palabra in indicadores_visual
        )

        if not tiene_indicios:
            return None

        prompt = f"""Analiza este texto de la sección "{titulo_seccion}" y propone un modelo visual útil.

TEXTO:
{contenido[:1500]}

REGLAS:
1. Si hay pasos, fases o secuencia, usa tipo "diagrama_flujo" con lista ordenada.
2. Si hay conceptos y relaciones, usa tipo "mapa_conceptual" con nodo central y nodos relacionados.
3. Mantén los textos de nodos/pasos cortos (máx. 6-8 palabras) y no más de 6 elementos.
4. Responde SOLO en formato JSON.
5. Si NO hay material visualizable, responde: {{"tiene_visual": false}}

Formato si hay visual:
{{
  "tiene_visual": true,
  "tipo_visual": "diagrama_flujo|mapa_conceptual",
  "titulo": "Título descriptivo",
  "elementos": ["Paso 1", "Paso 2", "Paso 3"],
  "nodo_central": "Concepto central",
  "nodos_relacionados": ["Nodo 1", "Nodo 2", "Nodo 3"]
}}

Responde SOLO el JSON, sin explicaciones.
"""

        try:
            response = await asyncio.to_thread(
                self.client.chat.complete,
                model=self.modelo,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.2,
                response_format={"type": "json_object"}
            )

            import json
            resultado = json.loads(response.choices[0].message.content)

            if not resultado.get('tiene_visual', False):
                return None

            return {
                'tipo': resultado.get('tipo_visual'),
                'titulo': resultado.get('titulo', 'Modelo visual'),
                'elementos': resultado.get('elementos', []),
                'nodo_central': resultado.get('nodo_central', ''),
                'nodos_relacionados': resultado.get('nodos_relacionados', [])
            }

        except Exception as e:
            logging.error(f"Error extrayendo visuales: {e}")
            return None


class GeneradorGraficos:
    """Genera gráficos profesionales con matplotlib."""
    
    @staticmethod
    def generar_grafico(datos_visual: Dict) -> BytesIO:
        """
        Genera un gráfico y lo devuelve como imagen en memoria.
        
        Args:
            datos_visual: Dict con 'tipo', 'titulo', 'datos', 'eje_x', 'eje_y'
        
        Returns:
            BytesIO con la imagen PNG
        """
        tipo = datos_visual['tipo']
        titulo = datos_visual['titulo']
        datos = datos_visual['datos']
        
        # Preparar datos
        etiquetas = list(datos.keys())
        valores = [float(v) for v in datos.values()]
        
        # Crear figura
        fig, ax = plt.subplots(figsize=(10, 6))
        
        if tipo == 'barras':
            barras = ax.bar(etiquetas, valores, color=COLORES_PROFESIONALES[:len(etiquetas)])
            
            # Añadir valores sobre las barras
            for barra in barras:
                altura = barra.get_height()
                ax.text(barra.get_x() + barra.get_width()/2., altura,
                       f'{altura:.1f}',
                       ha='center', va='bottom', fontweight='bold')
            
            ax.set_ylabel(datos_visual.get('eje_y', 'Valor'))
            if datos_visual.get('eje_x'):
                ax.set_xlabel(datos_visual['eje_x'])
            plt.xticks(rotation=45, ha='right')
        
        elif tipo == 'lineas':
            ax.plot(etiquetas, valores, marker='o', linewidth=2.5, 
                   markersize=8, color=COLORES_PROFESIONALES[0])
            
            # Añadir valores en los puntos
            for i, (x, y) in enumerate(zip(etiquetas, valores)):
                ax.text(i, y, f'{y:.1f}', ha='center', va='bottom', fontweight='bold')
            
            ax.set_ylabel(datos_visual.get('eje_y', 'Valor'))
            if datos_visual.get('eje_x'):
                ax.set_xlabel(datos_visual['eje_x'])
            ax.grid(True, alpha=0.3)
            plt.xticks(rotation=45, ha='right')
        
        elif tipo == 'tarta':
            # Crear gráfico de tarta
            colores = COLORES_PROFESIONALES[:len(etiquetas)]
            wedges, texts, autotexts = ax.pie(valores, labels=etiquetas, autopct='%1.1f%%',
                                              colors=colores, startangle=90)
            
            # Mejorar legibilidad
            for autotext in autotexts:
                autotext.set_color('white')
                autotext.set_fontweight('bold')
                autotext.set_fontsize(10)
            
            for text in texts:
                text.set_fontsize(11)
            
            ax.axis('equal')
        
        elif tipo == 'puntos':
            ax.scatter(range(len(valores)), valores, s=100, 
                      color=COLORES_PROFESIONALES[0], alpha=0.6, edgecolors='black')
            
            ax.set_ylabel(datos_visual.get('eje_y', 'Valor'))
            if datos_visual.get('eje_x'):
                ax.set_xlabel(datos_visual['eje_x'])
            ax.set_xticks(range(len(etiquetas)))
            ax.set_xticklabels(etiquetas, rotation=45, ha='right')
            ax.grid(True, alpha=0.3)
        
        # Título del gráfico
        plt.title(titulo, fontsize=14, fontweight='bold', pad=20)
        
        # Ajustar layout
        plt.tight_layout()
        
        # Guardar en memoria
        buffer = BytesIO()
        plt.savefig(buffer, format='png', dpi=150, bbox_inches='tight')
        plt.close(fig)
        
        buffer.seek(0)
        return buffer


class GeneradorVisuales:
    """Genera diagramas e ilustraciones conceptuales con matplotlib."""

    @staticmethod
    def _formatear_texto(texto: str, max_chars: int, ancho_linea: int) -> str:
        import textwrap

        texto = texto.strip()
        if len(texto) > max_chars:
            texto = textwrap.shorten(texto, width=max_chars, placeholder="…")
        return textwrap.fill(texto, width=ancho_linea)

    @staticmethod
    def _render_diagrama_flujo(titulo: str, elementos: List[str]) -> BytesIO:
        fig, ax = plt.subplots(figsize=(8, max(4, len(elementos) * 1.2)))
        ax.axis('off')

        box_width = 0.7
        box_height = 0.12
        start_y = 0.9
        step = 0.18

        for idx, texto in enumerate(elementos[:6]):
            y = start_y - idx * step
            rect = patches.FancyBboxPatch(
                (0.15, y - box_height / 2),
                box_width,
                box_height,
                boxstyle="round,pad=0.02,rounding_size=0.02",
                linewidth=1.5,
                edgecolor='#2E86C1',
                facecolor='#EAF2F8'
            )
            ax.add_patch(rect)
            texto_formateado = GeneradorVisuales._formatear_texto(texto, max_chars=60, ancho_linea=28)
            fontsize = 9 if len(texto_formateado) > 40 else 10
            ax.text(0.5, y, texto_formateado, ha='center', va='center', fontsize=fontsize, wrap=True)

            if idx < min(len(elementos), 6) - 1:
                ax.annotate(
                    '',
                    xy=(0.5, y - box_height / 2 - 0.02),
                    xytext=(0.5, y - step + box_height / 2 + 0.02),
                    arrowprops=dict(arrowstyle='->', color='#566573', lw=1.5)
                )

        ax.set_title(titulo, fontsize=13, fontweight='bold', pad=12)

        buffer = BytesIO()
        plt.tight_layout()
        plt.savefig(buffer, format='png', dpi=150, bbox_inches='tight')
        plt.close(fig)
        buffer.seek(0)
        return buffer

    @staticmethod
    def _render_mapa_conceptual(titulo: str, nodo_central: str, nodos: List[str]) -> BytesIO:
        fig, ax = plt.subplots(figsize=(8, 6))
        ax.axis('off')

        centro = (0.5, 0.5)
        ax.add_patch(patches.Circle(centro, 0.12, color='#D6EAF8', ec='#2E86C1', lw=2))
        nodo_central = GeneradorVisuales._formatear_texto(nodo_central or 'Concepto', max_chars=40, ancho_linea=14)
        ax.text(centro[0], centro[1], nodo_central, ha='center', va='center', fontsize=10)

        radio = 0.32
        nodos = nodos[:6]
        total = max(len(nodos), 1)
        for i, nodo in enumerate(nodos):
            angulo = 2 * np.pi * i / total
            x = centro[0] + radio * np.cos(angulo)
            y = centro[1] + radio * np.sin(angulo)
            ax.plot([centro[0], x], [centro[1], y], color='#7F8C8D', lw=1.2)
            ax.add_patch(patches.FancyBboxPatch(
                (x - 0.18, y - 0.05),
                0.36,
                0.1,
                boxstyle="round,pad=0.02,rounding_size=0.02",
                linewidth=1,
                edgecolor='#27AE60',
                facecolor='#E8F8F5'
            ))
            nodo_formateado = GeneradorVisuales._formatear_texto(nodo, max_chars=36, ancho_linea=16)
            fontsize = 8 if len(nodo_formateado) > 30 else 9
            ax.text(x, y, nodo_formateado, ha='center', va='center', fontsize=fontsize, wrap=True)

        ax.set_title(titulo, fontsize=13, fontweight='bold', pad=12)

        buffer = BytesIO()
        plt.tight_layout()
        plt.savefig(buffer, format='png', dpi=150, bbox_inches='tight')
        plt.close(fig)
        buffer.seek(0)
        return buffer

    def generar_visual(self, visual_spec: Dict) -> BytesIO:
        tipo = visual_spec.get('tipo')
        titulo = visual_spec.get('titulo', 'Modelo visual')

        if tipo == 'diagrama_flujo':
            elementos = visual_spec.get('elementos', [])
            if not elementos:
                elementos = ['Paso 1', 'Paso 2', 'Paso 3']
            return self._render_diagrama_flujo(titulo, elementos)

        if tipo == 'mapa_conceptual':
            nodo_central = visual_spec.get('nodo_central', 'Concepto')
            nodos = visual_spec.get('nodos_relacionados', [])
            return self._render_mapa_conceptual(titulo, nodo_central, nodos)

        elementos = visual_spec.get('elementos', ['Idea 1', 'Idea 2', 'Idea 3'])
        return self._render_diagrama_flujo(titulo, elementos)


class IntegradorGraficosWord:
    """Integra gráficos en documentos Word."""
    
    def __init__(self, client, modelo="mistral-large-latest"):
        self.extractor = ExtractorDatos(client, modelo)
        self.generador = GeneradorGraficos()
        self.extractor_visual = ExtractorVisuales(client, modelo)
        self.generador_visual = GeneradorVisuales()
    
    async def procesar_seccion_con_graficos(
        self, 
        titulo_seccion: str, 
        contenido: str, 
        doc,
        numero_seccion: int
    ) -> Dict[str, int]:
        """
        Procesa una sección, detecta datos, genera gráficos y los añade al doc.
        
        Args:
            titulo_seccion: Título de la sección
            contenido: Contenido de texto
            doc: Objeto Document de python-docx
            numero_seccion: Número de la sección
        
        Returns:
            Dict con conteos de recursos visuales añadidos
        """
        from docx.shared import Inches, Pt
        
        graficos_añadidos = 0
        visuales_añadidos = 0
        
        # Añadir título de sección
        doc.add_heading(titulo_seccion, level=1)
        
        # Dividir contenido en párrafos
        parrafos = contenido.split('\n\n')
        
        for i, parrafo in enumerate(parrafos):
            if not parrafo.strip():
                continue
            
            # Añadir párrafo
            doc.add_paragraph(parrafo.strip())
            
            # Cada 2-3 párrafos, intentar extraer datos y generar gráfico
            if i > 0 and i % 2 == 0 and graficos_añadidos < 2:  # Max 2 gráficos por sección
                try:
                    # Detectar y extraer datos del contexto acumulado
                    contexto = '\n'.join(parrafos[:i+1])
                    datos_visual = await self.extractor.detectar_y_extraer_datos(
                        contexto, 
                        titulo_seccion
                    )
                    
                    if datos_visual:
                        logging.info(f"✅ Datos encontrados en '{titulo_seccion}': {datos_visual['tipo']}")
                        
                        # Generar gráfico
                        imagen_buffer = self.generador.generar_grafico(datos_visual)
                        
                        # Insertar en Word
                        doc.add_paragraph()  # Espacio
                        doc.add_picture(imagen_buffer, width=Inches(5.5))
                        
                        # Centrar imagen
                        last_paragraph = doc.paragraphs[-1]
                        last_paragraph.alignment = 1  # WD_ALIGN_PARAGRAPH.CENTER
                        
                        # Añadir pie de figura
                        caption = doc.add_paragraph(
                            f"Figura {numero_seccion}.{graficos_añadidos + 1}: {datos_visual['titulo']}"
                        )
                        caption.alignment = 1
                        caption.runs[0].italic = True
                        caption.runs[0].font.size = Pt(10)
                        
                        doc.add_paragraph()  # Espacio después
                        graficos_añadidos += 1
                        
                        logging.info(f"✅ Gráfico {graficos_añadidos} añadido a sección {numero_seccion}")
                
                except Exception as e:
                    logging.error(f"Error generando gráfico en sección {numero_seccion}: {e}")
                    continue

            if i > 0 and i % 3 == 0 and visuales_añadidos < 1:
                try:
                    contexto = '\n'.join(parrafos[:i+1])
                    visual_spec = await self.extractor_visual.detectar_visual(
                        contexto,
                        titulo_seccion
                    )

                    if visual_spec:
                        logging.info(f"✅ Visual encontrado en '{titulo_seccion}': {visual_spec['tipo']}")
                        imagen_buffer = self.generador_visual.generar_visual(visual_spec)

                        doc.add_paragraph()
                        doc.add_picture(imagen_buffer, width=Inches(5.8))

                        last_paragraph = doc.paragraphs[-1]
                        last_paragraph.alignment = 1

                        caption = doc.add_paragraph(
                            f"Figura {numero_seccion}.{graficos_añadidos + visuales_añadidos + 1}: {visual_spec['titulo']}"
                        )
                        caption.alignment = 1
                        caption.runs[0].italic = True
                        caption.runs[0].font.size = Pt(10)

                        doc.add_paragraph()
                        visuales_añadidos += 1

                        logging.info(
                            f"✅ Visual {visuales_añadidos} añadido a sección {numero_seccion}"
                        )
                except Exception as e:
                    logging.error(f"Error generando visual en sección {numero_seccion}: {e}")
                    continue
        
        return {
            "graficos": graficos_añadidos,
            "visuales": visuales_añadidos,
            "total": graficos_añadidos + visuales_añadidos
        }


# Función auxiliar para usar desde brain_v2.py
async def añadir_graficos_inteligentes(
    doc, 
    secciones: List[Dict], 
    client, 
    context=None, 
    chat_id=None
):
    """
    Añade gráficos inteligentes a un documento Word basándose en el contenido.
    
    Args:
        doc: Documento Word (python-docx)
        secciones: Lista de dicts con 'titulo', 'contenido', 'numero'
        client: Cliente de Mistral
        context: Context de Telegram (opcional, para reporting)
        chat_id: ID de chat (opcional, para reporting)
    
    Returns:
        Número total de gráficos añadidos
    """
    from docx.shared import Pt
    
    integrador = IntegradorGraficosWord(client)
    total_graficos = 0
    
    for seccion in secciones:
        try:
            if context and chat_id:
                await context.bot.send_message(
                    chat_id=chat_id,
                    text=f"🔍 Buscando datos visualizables en: {seccion['titulo'][:50]}..."
                )
            
            resultado = await integrador.procesar_seccion_con_graficos(
                titulo_seccion=seccion['titulo'],
                contenido=seccion['contenido'],
                doc=doc,
                numero_seccion=seccion['numero']
            )
            
            total_graficos += resultado["total"]
            
            if resultado["total"] > 0 and context and chat_id:
                await context.bot.send_message(
                    chat_id=chat_id,
                    text=f"✅ {resultado['total']} recurso(s) visual(es) añadido(s) a la sección {seccion['numero']}"
                )
        
        except Exception as e:
            logging.error(f"Error procesando sección {seccion.get('numero', '?')}: {e}")
            continue
    
    return total_graficos
