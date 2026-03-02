"""
BRAIN V2 - Sistema de procesamiento mejorado con motores especializados
========================================================================
Cambios principales:
- IntegraciГіn de ContentEngine para generaciГіn de estudios
- SeparaciГіn clara entre conversaciГіn y generaciГіn de contenido
- Mejor manejo de contexto y memoria
- ValidaciГіn de resultados
"""

import os
import json
import logging
import re
import asyncio
import unicodedata
import subprocess
import sys
import requests
from datetime import datetime, timedelta
from zoneinfo import ZoneInfo
from datetime import time
from docx import Document
from telegram import Update, ReplyKeyboardMarkup
from telegram.ext import ContextTypes
from mistralai import Mistral

# Importar motores especializados
from content_engine import ContentEngine, AnalisisFinanciero

# Importar herramientas existentes
from tools_finance import (
    analizar_inversiones,
    buscar_oportunidades_inversion,
    obtener_lista_seguimiento
)
from tools_system import (
    crear_evento_calendar, control_openhab, 
    obtener_ip_publica, buscar_internet, modelo_whisper, exportar_a_word_premium
)

from evaluador_profesional import EvaluadorProfesionalCartera, formatear_informe_profesional

# Estados de conversaciГіn
esperando_empresa_deep = {}
esperando_mercado_oportunidades = {}
ESPERANDO_PROMPT_STUDIO = {}

STUDIO_TIPOS = [
    {
        "id": "practica_asignatura",
        "label": "PrГЎctica de asignatura (ESO/FP InformГЎtica)",
        "keywords": ["practica", "prГЎctica", "fp", "informatica", "informГЎtica", "asignatura"],
    },
    {
        "id": "estudio_academico",
        "label": "Estudio acadГ©mico",
        "keywords": ["estudio academico", "academico", "acadГ©mico"],
    },
    {
        "id": "informe_profesional",
        "label": "Informe profesional",
        "keywords": ["informe profesional", "profesional"],
    },
    {
        "id": "reporte_tecnico",
        "label": "Reporte tГ©cnico",
        "keywords": ["reporte tecnico", "tГ©cnico", "tecnico"],
    },
    {
        "id": "resumen_ejecutivo",
        "label": "Resumen ejecutivo",
        "keywords": ["resumen ejecutivo", "ejecutivo", "resumen"],
    },
    {
        "id": "investigacion_mercado",
        "label": "InvestigaciГіn de mercado",
        "keywords": ["investigacion de mercado", "investigaciГіn de mercado", "mercado"],
    },
]

STUDIO_TONOS = [
    {"id": "formal", "label": "Formal", "keywords": ["formal"]},
    {"id": "academico", "label": "AcadГ©mico", "keywords": ["academico", "acadГ©mico"]},
    {"id": "informativo", "label": "Informativo", "keywords": ["informativo"]},
    {"id": "divulgativo", "label": "Divulgativo", "keywords": ["divulgativo"]},
    {"id": "creativo", "label": "Creativo", "keywords": ["creativo"]},
    {"id": "persuasivo", "label": "Persuasivo", "keywords": ["persuasivo"]},
]

STUDIO_EXTENSION = [
    {"id": "corto", "label": "Corto (1-2 pГЎginas)", "keywords": ["corto", "breve"]},
    {"id": "medio", "label": "Medio (3-5 pГЎginas)", "keywords": ["medio", "intermedio"]},
    {"id": "largo", "label": "Largo (6-10 pГЎginas)", "keywords": ["largo"]},
    {"id": "extenso", "label": "Extenso (+10 pГЎginas)", "keywords": ["extenso", "muy largo"]},
]

STUDIO_NIVEL = [
    {"id": "introductorio", "label": "Introductorio", "keywords": ["introductorio", "basico", "bГЎsico"]},
    {"id": "intermedio", "label": "Intermedio", "keywords": ["intermedio", "medio"]},
    {"id": "avanzado", "label": "Avanzado", "keywords": ["avanzado", "experto"]},
]



def _normalizar_texto(texto: str) -> str:
    texto = texto.strip().lower()
    return "".join(
        caracter
        for caracter in unicodedata.normalize("NFD", texto)
        if unicodedata.category(caracter) != "Mn"
    )


def _seleccionar_opcion(texto: str, opciones: list) -> dict | None:
    texto_norm = _normalizar_texto(texto)
    match = re.match(r"^(\d+)", texto_norm)
    if match:
        indice = int(match.group(1)) - 1
        if 0 <= indice < len(opciones):
            return opciones[indice]

    for opcion in opciones:
        if any(keyword in texto_norm for keyword in opcion["keywords"]):
            return opcion

    for opcion in opciones:
        if _normalizar_texto(opcion["label"]) == texto_norm:
            return opcion

    return None

# Modelos
MODELO_CONVERSACION = "mistral-small-latest"  # Para chat normal
MODELO_GENERACION = "mistral-large-latest"    # Para contenido complejo


def _chat_autorizado(config: dict, chat_id: int) -> bool:
    telegram_cfg = (config or {}).get("telegram", {})
    allowed_users = telegram_cfg.get("allowed_users", [])
    allowed_restricted_users = telegram_cfg.get("allowed_restricted_users", [])
    return chat_id in allowed_users or chat_id in allowed_restricted_users

# Historial de conversaciones
MEMORIA_PATH = os.getenv("ROBI_MEMORY_FILE", "robi_memoria.json")
historiales = {}


def _cargar_memoria_persistente() -> dict:
    """Carga historiales guardados en disco si existen."""
    if not os.path.exists(MEMORIA_PATH):
        return {}

    try:
        with open(MEMORIA_PATH, "r", encoding="utf-8") as memoria_file:
            memoria = json.load(memoria_file)
        if not isinstance(memoria, dict):
            logging.warning("вљ пёЏ Memoria invГЎlida: formato no reconocido")
            return {}

        memoria_limpia = {}
        for user_id, mensajes in memoria.items():
            if isinstance(user_id, str) and isinstance(mensajes, list):
                memoria_limpia[user_id] = mensajes

        if memoria_limpia:
            logging.info(f"рџ§  Memoria cargada: {len(memoria_limpia)} conversaciones")
        return memoria_limpia
    except Exception as e:
        logging.warning(f"вљ пёЏ No se pudo cargar memoria persistente: {e}")
        return {}


def _guardar_memoria_persistente() -> None:
    """Guarda historiales en disco para mantener contexto entre reinicios."""
    try:
        directorio = os.path.dirname(MEMORIA_PATH)
        if directorio:
            os.makedirs(directorio, exist_ok=True)

        with open(MEMORIA_PATH, "w", encoding="utf-8") as memoria_file:
            json.dump(historiales, memoria_file, ensure_ascii=False, indent=2)
    except Exception as e:
        logging.warning(f"вљ пёЏ No se pudo guardar memoria persistente: {e}")


historiales = _cargar_memoria_persistente()


def generar_prompt_sistema():
    """Genera el prompt del sistema con fecha actualizada."""
    hoy = datetime.now()
    fecha_str = hoy.strftime("%Y-%m-%d")
    dia_semana = ["lunes", "martes", "miГ©rcoles", "jueves", "viernes", "sГЎbado", "domingo"][hoy.weekday()]
    
    return f"""Eres Robi, un asistente domГ©stico inteligente especializado.

FECHA ACTUAL: {fecha_str} ({dia_semana})

=== COMANDOS DISPONIBLES ===

1пёЏвѓЈ BГљSQUEDA WEB
Formato: BUSCAR: 'consulta'
Ejemplo: BUSCAR: 'noticias tecnologГ­a IA'

2пёЏвѓЈ DOMГ“TICA
Formato: ACCION: 'dispositivo', 'ON'/'OFF'
Dispositivos: salon, despacho, cocina, comedor, dormitorio, caterina, ovidi, 
             ventilador despacho, ventilador caterina, ventilador ovidi, ventilador dormitori

3пёЏвѓЈ CALENDARIO
Formato: CALENDAR_CREAR: 'tГ­tulo', 'YYYY-MM-DD', 'HH:MM'
IMPORTANTE: Calcula fechas relativas desde {fecha_str}
- "maГ±ana" = {(hoy + timedelta(days=1)).strftime("%Y-%m-%d")}
- "pasado maГ±ana" = {(hoy + timedelta(days=2)).strftime("%Y-%m-%d")}

4пёЏвѓЈ INVERSIONES
- Ver cartera: CONSULTAR: 'INVERSIONES'
- AnГЎlisis valor: ANALIZAR_VALOR: 'nombre_empresa'
- InvestigaciГіn profunda: DEEP_RESEARCH: 'ticker'
- Buscar oportunidades: BUSCAR_OPORTUNIDADES: 'nasdaq' o 'ibex'
- Evaluar cartera completa: EVALUAR_CARTER

5пёЏвѓЈ UTILIDADES
- IP pГєblica: CONSULTAR: 'IP'

=== REGLAS CRГЌTICAS ===
- Responde de forma CONCISA y DIRECTA
- NO uses comandos para preguntas que puedes responder tГє
- USA comandos solo cuando sea necesario obtener datos externos
- Para anГЎlisis financiero, SIEMPRE busca datos actuales primero
- Formato de comandos EXACTO, sin variaciones
"""


async def handle_text(update: Update, context: ContextTypes.DEFAULT_TYPE, client, config, texto_override=None):
    """Handler principal de mensajes de texto."""
    
    # Control de acceso
    if not _chat_autorizado(config, update.effective_chat.id):
        logging.warning(f"Usuario no autorizado: {update.effective_chat.id}")
        return
    
    user_text = texto_override or update.message.text
    user_id = update.effective_chat.id
    
    # Revisar si estГЎ esperando input para Studio
    if user_id in ESPERANDO_PROMPT_STUDIO:
        # Recuperar client y config del context.user_data
        client_studio = context.user_data.get('client', client)
        config_studio = context.user_data.get('config', config)
        await recibir_prompt_studio(update, context, client_studio, config_studio)
        return
    
    # Revisar si estГЎ esperando input para Deep Research
    if user_id in esperando_empresa_deep and esperando_empresa_deep[user_id]:
        esperando_empresa_deep[user_id] = False
        await update.message.reply_text(
            f"рџљЂ Evaluando {user_text.upper()} con anГЎlisis profesional...\n\n"
            "Esto llevarГЎ unos minutos. Te aviso cuando termine."
        )
        context.application.create_task(ejecutar_evaluacion_valor(update, user_text, client, config))
        return
    
    # Revisar si estГЎ esperando mercado para oportunidades
    if user_id in esperando_mercado_oportunidades and esperando_mercado_oportunidades[user_id]:
        esperando_mercado_oportunidades[user_id] = False
        await lanzar_escaner_oportunidades(update, user_text, client, config)
        return
    
    # Procesar mensaje normal
    await handle_message_logic(update, context, user_text, client, config)


async def handle_message_logic(update, context, user_text, client, config, retorno_texto=False):
    """
    LГіgica principal de procesamiento de mensajes.
    
    Args:
        update: Update de Telegram (puede ser None para llamadas internas)
        context: Contexto de Telegram
        user_text: Texto del usuario
        client: Cliente de Mistral
        config: ConfiguraciГіn
        retorno_texto: Si True, devuelve texto en lugar de enviarlo
    """
    
    user_id = str(update.effective_chat.id) if update else "SYSTEM_AGENT"
    
    # TraducciГіn de botones
    if user_text == "рџ“€ Mi Cartera":
        user_text = "Dime cГіmo van mis inversiones"
    elif user_text == "рџ”Ќ Buscar Oportunidades":
        user_text = "Busca oportunidades de inversiГіn en el Nasdaq e Ibex"
    elif user_text == "рџЊђ Mi IP":
        user_text = "Dime mi IP pГєblica"
    
    # Inicializar historial
    prompt_sistema = generar_prompt_sistema()
    if user_id not in historiales:
        historiales[user_id] = [
            {"role": "system", "content": prompt_sistema},
            {"role": "assistant", "content": "Robi operativo. ВїEn quГ© puedo ayudarte?"}
        ]
    else:
        # Actualizar prompt del sistema con fecha actual
        historiales[user_id][0] = {"role": "system", "content": prompt_sistema}
    
    # AГ±adir mensaje del usuario
    historiales[user_id].append({"role": "user", "content": user_text})
    _guardar_memoria_persistente()
    
    # Recortar historial (mantener Гєltimo 10 + sistema)
    if len(historiales[user_id]) > 11:
        historiales[user_id] = [historiales[user_id][0]] + historiales[user_id][-10:]
    
    try:
        # Llamada a Mistral
        chat_response = client.chat.complete(
            model=MODELO_CONVERSACION,
            messages=historiales[user_id],
            temperature=0.4
        )
        texto_ai = chat_response.choices[0].message.content
        
        logging.info(f"рџ¤– Respuesta: {texto_ai}")
        
        # Guardar respuesta en historial
        historiales[user_id].append({"role": "assistant", "content": texto_ai})
        _guardar_memoria_persistente()
        
        # Procesar comandos
        respuesta_final = await procesar_comandos(texto_ai, client, config)
        
        # Devolver o enviar
        if retorno_texto:
            return respuesta_final
        else:
            await enviar_mensaje_largo(update, respuesta_final)
    
    except Exception as e:
        logging.error(f"Error en handle_message_logic: {e}")
        mensaje_error = f"вќЊ Error: {str(e)}"
        
        if retorno_texto:
            return mensaje_error
        else:
            await update.message.reply_text(mensaje_error)


async def procesar_comandos(texto_ai: str, client, config) -> str:
    """Procesa comandos detectados en la respuesta de la IA."""
    
    respuesta = texto_ai
    
    # BUSCAR
    if "BUSCAR:" in texto_ai:
        match = re.search(r"BUSCAR:\s*['\"](.+?)['\"]", texto_ai)
        if match:
            query = match.group(1)
            logging.info(f"рџ”Ќ Ejecutando bГєsqueda: {query}")
            resultado = await asyncio.to_thread(
                buscar_internet, query, client, config, MODELO_GENERACION
            )
            respuesta = texto_ai.replace(match.group(0), resultado)
    
    # ACCION (domГіtica)
    if "ACCION:" in texto_ai:
        match = re.search(r"ACCION:\s*['\"](.+?)['\"]\s*,\s*['\"](.+?)['\"]", texto_ai)
        if match:
            item, state = match.group(1), match.group(2)
            logging.info(f"рџЏ  AcciГіn domГіtica: {item} -> {state}")
            
            resultado = await asyncio.to_thread(control_openhab, item, state, config)
            
            respuesta = texto_ai.replace(match.group(0), resultado)
    
    # CALENDAR_CREAR
    if "CALENDAR_CREAR:" in texto_ai:
        match = re.search(r"CALENDAR_CREAR:\s*['\"](.+?)['\"]\s*,\s*['\"](.+?)['\"]\s*,\s*['\"](.+?)['\"]", texto_ai)
        if match:
            titulo, fecha, hora = match.group(1), match.group(2), match.group(3)
            logging.info(f"рџ“… Creando evento: {titulo} - {fecha} {hora}")
            resultado = await asyncio.to_thread(crear_evento_calendar, titulo, fecha, hora)
            respuesta = texto_ai.replace(match.group(0), resultado)
    
    # CONSULTAR: INVERSIONES
    if "CONSULTAR: 'INVERSIONES'" in texto_ai or "CONSULTAR: \"INVERSIONES\"" in texto_ai:
        logging.info("рџ“Љ Consultando inversiones...")
        resultado = await asyncio.to_thread(analizar_inversiones)
        respuesta = texto_ai.replace("CONSULTAR: 'INVERSIONES'", resultado)
        respuesta = respuesta.replace("CONSULTAR: \"INVERSIONES\"", resultado)
    
    # EVALUAR_CARTERA
    if "EVALUAR_CARTERA" in texto_ai:
        logging.info("рџ”Ќ Ejecutando evaluaciГіn de cartera...")
        # Nota: Esto se ejecuta sГ­ncronamente en el hilo de IA
        # Para mejor UX, el usuario deberГ­a usar /evaluar directamente
        from tools_finance import analizar_inversiones
        
        # Primero verificar que hay Excel
        ruta_excel = "/app/documentos/bolsav2.xlsx"
        if not os.path.exists(ruta_excel):
            respuesta = texto_ai.replace("EVALUAR_CARTERA", 
                "вќЊ No encuentro tu archivo de inversiones.")
        else:
            respuesta = texto_ai.replace("EVALUAR_CARTERA",
                "Para evaluar tu cartera completa, usa el comando /evaluar. "
                "El anГЎlisis llevarГЎ unos minutos pero recibirГЎs recomendaciones "
                "detalladas para cada valor.")
        
        return respuesta


    # CONSULTAR: IP
    if "CONSULTAR: 'IP'" in texto_ai or "CONSULTAR: \"IP\"" in texto_ai:
        logging.info("рџЊђ Obteniendo IP...")
        resultado = await asyncio.to_thread(obtener_ip_publica)
        respuesta = texto_ai.replace("CONSULTAR: 'IP'", resultado)
        respuesta = respuesta.replace("CONSULTAR: \"IP\"", resultado)
    
    # ANALIZAR_VALOR
    if "ANALIZAR_VALOR:" in texto_ai:
        match = re.search(r"ANALIZAR_VALOR:\s*['\"](.+?)['\"]", texto_ai)
        if match:
            valor = match.group(1)
            logging.info(f"рџ“€ Analizando valor: {valor}")
            # AquГ­ podrГ­as usar el nuevo AnalisisFinanciero si quieres
            # Por ahora mantenemos la funciГіn original
            from tools_finance import super_asesor_financiero
            resultado = await asyncio.to_thread(
                super_asesor_financiero,
                valor,
                client,
                lambda q: buscar_internet(q, client, config, MODELO_GENERACION)
            )
            respuesta = texto_ai.replace(match.group(0), resultado)
    
    return respuesta


async def handle_voice(update: Update, context: ContextTypes.DEFAULT_TYPE, client, config):
    """Handler de mensajes de voz."""
    
    if not _chat_autorizado(config, update.effective_chat.id):
        return
    
    try:
        # Descargar archivo de voz
        voice_file = await update.message.voice.get_file()
        voice_path = "temp_voice.ogg"
        await voice_file.download_to_drive(voice_path)
        
        # Transcribir
        result = modelo_whisper.transcribe(voice_path)
        texto_transcrito = result["text"]
        
        logging.info(f"рџЋ¤ Voz transcrita: {texto_transcrito}")
        
        # Procesar como texto normal
        await handle_text(update, context, client, config, texto_override=texto_transcrito)
        
        # Limpiar archivo temporal
        if os.path.exists(voice_path):
            os.remove(voice_path)
    
    except Exception as e:
        logging.error(f"Error procesando voz: {e}")
        await update.message.reply_text(f"вќЊ Error al procesar audio: {str(e)}")


async def enviar_mensaje_largo(update: Update, texto: str):
    """EnvГ­a mensajes largos dividiГ©ndolos si es necesario."""
    MAX_LENGTH = 4000
    
    try:
        if len(texto) > MAX_LENGTH:
            for i in range(0, len(texto), MAX_LENGTH):
                await update.message.reply_text(texto[i:i+MAX_LENGTH])
        else:
            await update.message.reply_text(texto)
    except Exception as e:
        logging.error(f"Error enviando mensaje: {e}")


async def enviar_mensaje_largo_chat(context: ContextTypes.DEFAULT_TYPE, chat_id: int, texto: str):
    """EnvГ­a mensajes largos a un chat sin depender del objeto update."""
    MAX_LENGTH = 4000
    try:
        if len(texto) > MAX_LENGTH:
            for i in range(0, len(texto), MAX_LENGTH):
                await context.bot.send_message(chat_id=chat_id, text=texto[i:i + MAX_LENGTH])
        else:
            await context.bot.send_message(chat_id=chat_id, text=texto)
    except Exception as e:
        logging.error(f"Error enviando mensaje largo al chat {chat_id}: {e}")


def _normalizar_informe_para_telegram_movil(texto: str) -> str:
    """Adapta el informe para que se lea bien en Telegram mГіvil."""
    if not texto:
        return ""

    # Evitar formatos que se rompen en pantallas pequeГ±as.
    texto = texto.replace("```", "")
    texto = re.sub(r"^#{1,6}\s*", "", texto, flags=re.MULTILINE)
    texto = texto.replace("**", "").replace("__", "")

    # Convertir viГ±etas markdown a un formato uniforme y simple.
    texto = re.sub(r"^\s*[-*]\s+", "вЂў ", texto, flags=re.MULTILINE)

    # Quitar tablas markdown (quedan descuadradas en mГіvil) y dejarlas como lГ­neas separadas.
    lineas_limpias = []
    for linea in texto.splitlines():
        if "|" in linea:
            partes = [p.strip() for p in linea.split("|") if p.strip()]
            if not partes:
                continue
            if all(set(p) <= {":", "-"} for p in partes):
                continue
            linea = " вЂў ".join(partes)
        lineas_limpias.append(linea.rstrip())

    texto = "\n".join(lineas_limpias)

    # Normalizar saltos y espacios.
    texto = re.sub(r"\n{3,}", "\n\n", texto)
    texto = re.sub(r"[ \t]{2,}", " ", texto)

    return texto.strip()


def _buscar_contexto_mercados_diario(config: dict, query: str) -> str:
    """Consulta Tavily sin timeout explГ­cito para no cortar informes largos."""
    api_key = config.get("tavily", {}).get("api_key")
    if not api_key:
        return "вљ пёЏ No hay API key de Tavily configurada."

    response = requests.post(
        "https://api.tavily.com/search",
        json={
            "api_key": api_key,
            "query": query,
            "search_depth": "advanced",
            "max_results": 8,
            "include_answer": True,
        },
    )
    response.raise_for_status()
    data = response.json()
    resultados = data.get("results", [])

    bloques = []
    answer = data.get("answer")
    if answer:
        bloques.append(f"Resumen inicial: {answer}")

    for i, item in enumerate(resultados, start=1):
        titulo = item.get("title", "Sin tГ­tulo")
        contenido = item.get("content", "Sin contenido")
        url = item.get("url", "")
        bloques.append(f"Fuente {i}: {titulo}\n{contenido}\nURL: {url}")

    return "\n\n".join(bloques) if bloques else "Sin resultados relevantes."


async def generar_informe_studio_diario(client, config) -> str:
    """Genera el informe diario de mercados y cartera (sin lГ­mite interno de tiempo)."""
    fecha_hoy = datetime.now(ZoneInfo("Europe/Madrid"))
    fecha_ayer = fecha_hoy - timedelta(days=1)
    resumen_cartera = await asyncio.to_thread(analizar_inversiones)

    query = (
        "latest market recap previous session Nasdaq Composite IBEX 35 gold spot XAUUSD "
        "main drivers why markets moved analyst commentary macro events and what to watch today"
    )
    contexto_mercados = await asyncio.to_thread(_buscar_contexto_mercados_diario, config, query)

    prompt = f"""
    Eres un analista financiero senior. Genera un informe diario en espaГ±ol, claro y accionable.

    FECHA INFORME: {fecha_hoy.strftime('%Y-%m-%d')}
    DГЌA ANALIZADO (AYER): {fecha_ayer.strftime('%Y-%m-%d')}

    DATOS DE CARTERA DEL USUARIO:
    {resumen_cartera}

    CONTEXTO DE MERCADOS (WEB):
    {contexto_mercados}

    OBJETIVO DEL INFORME (OBLIGATORIO):
    1) QuГ© ocurriГі ayer con sus acciones.
    2) Por quГ© subieron o cayeron (motivos concretos si estГЎn disponibles).
    3) QuГ© pasГі en Nasdaq, IBEX y oro.
    4) QuГ© se espera para hoy (escenarios y catalizadores a vigilar).

    FORMATO DE SALIDA (OPTIMIZADO PARA TELEGRAM MГ“VIL):
    - TГ­tulo: "рџ“Љ Studio Diario - Resumen de Mercados"
    - SecciГіn 1: "Tu cartera ayer"
    - SecciГіn 2: "Mercados: Nasdaq, IBEX y Oro"
    - SecciГіn 3: "Claves para hoy"
    - SecciГіn 4: "Plan rГЎpido (3 acciones recomendadas)"
    - Estilo claro, sin relleno, y con viГ±etas.
    - NO uses tablas, NO uses markdown complejo, NO uses bloques de cГіdigo.
    - Escribe frases cortas (mГЎximo 1-2 lГ­neas por viГ±eta en mГіvil).
    - Deja una lГ­nea en blanco entre secciones para legibilidad.
    """

    respuesta = await asyncio.to_thread(
        client.chat.complete,
        model=MODELO_GENERACION,
        messages=[{"role": "user", "content": prompt}],
        temperature=0.2,
    )
    contenido = respuesta.choices[0].message.content
    return _normalizar_informe_para_telegram_movil(contenido)


async def ejecutar_studio_diario_en_background(chat_id: int, context: ContextTypes.DEFAULT_TYPE, client, config):
    """Ejecuta Studio Diario en background y envГ­a el resultado al chat indicado."""
    try:
        informe = await generar_informe_studio_diario(client, config)
        await enviar_mensaje_largo_chat(context, chat_id, informe)
    except Exception as e:
        logging.error(f"Error generando Studio Diario: {e}")
        await context.bot.send_message(
            chat_id=chat_id,
            text=(
                "вќЊ Error generando /studiodiario.\n"
                f"Detalle: {str(e)}"
            ),
        )


async def studio_diario_programado_callback(context: ContextTypes.DEFAULT_TYPE):
    """Job programado de lunes a viernes a las 07:00."""
    job_data = context.job.data or {}
    chat_id = job_data.get("chat_id")
    client = job_data.get("client")
    config = job_data.get("config")

    if not chat_id or not client or not config:
        logging.error("вќЊ Job Studio Diario sin parГЎmetros obligatorios")
        return

    await context.bot.send_message(
        chat_id=chat_id,
        text=(
            "вЏі Iniciando informe automГЎtico /studiodiario...\n"
            "Se estГЎ ejecutando en background y puede tardar lo necesario."
        ),
    )
    context.application.create_task(ejecutar_studio_diario_en_background(chat_id, context, client, config))


def programar_studio_diario(app, client, config):
    """Programa el informe diario de lunes a viernes a las 07:00 (Europe/Madrid)."""
    allowed_users = config.get("telegram", {}).get("allowed_users", [])
    if not allowed_users:
        logging.warning("вљ пёЏ No hay allowed_users para programar /studiodiario")
        return

    chat_id = allowed_users[0]
    tz = ZoneInfo("Europe/Madrid")
    hora_objetivo = time(hour=7, minute=0, tzinfo=tz)

    for job in app.job_queue.get_jobs_by_name("studio_diario_weekdays"):
        job.schedule_removal()

    app.job_queue.run_daily(
        studio_diario_programado_callback,
        time=hora_objetivo,
        days=(0, 1, 2, 3, 4),
        name="studio_diario_weekdays",
        data={"chat_id": chat_id, "client": client, "config": config},
    )
    logging.info("вњ… /studiodiario programado L-V 07:00 (Europe/Madrid)")


# ==================== COMANDOS ====================

async def start_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Comando /start."""
    user_id = str(update.effective_chat.id)
    
    # Limpiar historial
    if user_id in historiales:
        del historiales[user_id]
        _guardar_memoria_persistente()

    await update.message.reply_text("рџ”„ Actualizando Robi desde Git...")

    comando_pull = ["git", "pull", "--ff-only"]
    resultado_pull = await asyncio.to_thread(
        subprocess.run,
        comando_pull,
        cwd=os.path.dirname(os.path.abspath(__file__)),
        capture_output=True,
        text=True,
        timeout=60,
        check=False,
    )

    salida_pull = (resultado_pull.stdout or resultado_pull.stderr or "Sin salida").strip()
    if len(salida_pull) > 1200:
        salida_pull = f"{salida_pull[:1200]}..."

    if resultado_pull.returncode != 0:
        await update.message.reply_text(
            "вќЊ No se pudo actualizar el cГіdigo con git pull.\n\n"
            f"Salida:\n{salida_pull}"
        )
        await mostrar_menu_principal(update)
        return

    await update.message.reply_text(
        "вњ… CГіdigo actualizado correctamente.\n"
        "в™»пёЏ Reiniciando servidor de Robi para cargar cambios..."
    )

    context.application.create_task(_reiniciar_proceso_robicamente())


async def _reiniciar_proceso_robicamente():
    """Reinicia el proceso actual para cargar cГіdigo actualizado."""
    await asyncio.sleep(1)
    os.execv(sys.executable, [sys.executable, *sys.argv])


async def mostrar_menu_principal(update: Update):
    """Muestra el menГє principal con botones."""
    botones = [
        ['рџ“€ Mi Cartera', 'рџ”Ќ Buscar Oportunidades'],
        ['рџЊђ Mi IP', 'рџЏ  DomГіtica'],
        ['рџ“‹ Deep Research (Manual)']
    ]
    teclado = ReplyKeyboardMarkup(botones, resize_keyboard=True)
    
    await update.message.reply_text(
        "Robi reiniciado. ВїQuГ© necesitas?",
        reply_markup=teclado
    )


async def handle_command(update: Update, context: ContextTypes.DEFAULT_TYPE, client, config):
    """Handler de comandos especiales."""
    
    user_text = update.message.text
    
    if not _chat_autorizado(config, update.effective_chat.id):
        logging.warning(f"Usuario no autorizado (comando): {update.effective_chat.id}")
        return
    
    # DEEP RESEARCH (ejecutar en background)
    if user_text.startswith("/deep"):
        valor = user_text.replace("/deep", "").strip()
        
        if not valor:
            esperando_empresa_deep[update.effective_chat.id] = True
            await update.message.reply_text("рџ”Ќ ВїQuГ© empresa quieres investigar?")
            return
        
        await update.message.reply_text(
            f"рџљЂ Evaluando {valor.upper()} con anГЎlisis profesional...\n\n"
            "Esto llevarГЎ unos minutos. Te aviso cuando termine."
        )
        # CRГЌTICO: context.application.create_task para que se ejecute en background
        context.application.create_task(ejecutar_evaluacion_valor(update, valor, client, config))
        return
    
    # INVERSIONES (ejecutar directamente sin pasar por IA)
    elif user_text.startswith("/inversiones"):
        logging.info("рџ“Љ Comando /inversiones ejecutado directamente")
        try:
            resultado = await asyncio.to_thread(analizar_inversiones)
            await enviar_mensaje_largo(update, resultado)
        except Exception as e:
            logging.error(f"Error en /inversiones: {e}")
            await update.message.reply_text(f"вќЊ Error al consultar inversiones: {str(e)}")
        return

    # SEGUIMIENTO (lista watchlist V/X/Y)
    elif user_text.startswith("/seguimiento"):
        logging.info("рџ‘Ђ Comando /seguimiento ejecutado directamente")
        try:
            resultado = await asyncio.to_thread(obtener_lista_seguimiento)
            await enviar_mensaje_largo(update, resultado)
        except Exception as e:
            logging.error(f"Error en /seguimiento: {e}")
            await update.message.reply_text(f"вќЊ Error al consultar seguimiento: {str(e)}")
        return
    
    # EVALUAR CARTERA (ejecutar directamente)
    elif user_text.startswith("/evaluar"):
        logging.info("рџ”Ќ Comando /evaluar ejecutado - Evaluando cartera completa")
        await update.message.reply_text(
            "рџ”Ќ **Evaluando tu cartera completa...**\n\n"
            "Esto incluye:\n"
            "вЂў AnГЎlisis de cada posiciГіn\n"
            "вЂў Recomendaciones de compra/venta\n"
            "вЂў Precios objetivo\n"
            "вЂў Niveles de stop loss y take profit\n\n"
            "вЏ±пёЏ Esto llevarГЎ 2-5 minutos (depende del nГєmero de valores). "
            "Puedes seguir usando el bot mientras tanto."
        )
        # Ejecutar en background
        context.application.create_task(
            ejecutar_evaluacion_cartera(update, client, config),
            update=update
        )
        return

    # IP (ejecutar directamente sin pasar por IA)
    elif user_text.startswith("/ip"):
        logging.info("рџЊђ Comando /ip ejecutado directamente")
        resultado = await asyncio.to_thread(obtener_ip_publica)
        await update.message.reply_text(resultado)
        return
    
    # STUDIO DIARIO (ejecutar en background)
    elif user_text.startswith("/studiodiario"):
        await crear_studiodiario_command(update, context, client, config)
        return

    # OPORTUNIDADES (ejecutar en background)
    elif user_text.startswith("/oportunidades"):
        mercado = user_text.replace("/oportunidades", "").strip()
        
        if not mercado:
            esperando_mercado_oportunidades[update.effective_chat.id] = True
            botones = [['NASDAQ', 'IBEX'], ['CRYPTO', 'NYSE']]
            teclado = ReplyKeyboardMarkup(botones, one_time_keyboard=True, resize_keyboard=True)
            
            await update.message.reply_text(
                "рџ”Ќ ВїEn quГ© mercado busco?",
                reply_markup=teclado
            )
            return
        
        await update.message.reply_text(f"рџ”Ќ Escaneando el mercado {mercado.upper()}...\n\nEsto llevarГЎ un momento. Te aviso cuando encuentre algo.")
        # CRГЌTICO: context.application.create_task para background
        context.application.create_task(lanzar_escaner_oportunidades(update, mercado, client, config))
        return
    
    # Resto de comandos: NO pasar por memoria/IA.
    # Si llega aquГ­, el comando no estГЎ soportado y no debemos responder con contexto previo.
    await update.message.reply_text(
        "вќ“ Comando no reconocido. Usa /inversiones, /seguimiento, /evaluar, /oportunidades, /deep, /ip, /studio o /studiodiario."
    )


async def configurar_comandos(app):
    """Configura el menГє de comandos de Telegram."""
    comandos = [
        ("inversiones", "Ver balance de cartera"),
        ("seguimiento", "Ver lista de seguimiento"),
        ("evaluar", "Evaluar cartera y obtener recomendaciones"),
        ("oportunidades", "Buscar oportunidades de inversiГіn"),
        ("deep", "AnГЎlisis profundo de valor"),
        ("ip", "Consultar IP pГєblica"),
        ("studio", "Generar informe/estudio"),
        ("studiodiario", "Informe diario de mercados"),
        ("generarpartitura", "рџЋµ Generar lectura a vista para fagot"),
        ("start", "Reiniciar Robi")
    ]
    await app.bot.set_my_commands(comandos)



async def crear_studiodiario_command(update: Update, context: ContextTypes.DEFAULT_TYPE, client=None, config=None):
    """Comando /studiodiario - genera informe diario en background a demanda."""
    if not _chat_autorizado(config or {}, update.effective_chat.id):
        logging.warning(f"Usuario no autorizado (/studiodiario): {update.effective_chat.id}")
        return

    chat_id = update.effective_chat.id

    await update.message.reply_text(
        "рџљЂ /studiodiario iniciado en background.\n"
        "No tiene lГ­mite de tiempo y tardarГЎ lo necesario. Te enviarГ© el informe al terminar."
    )
    context.application.create_task(ejecutar_studio_diario_en_background(chat_id, context, client, config))



# ==================== STUDIO (NUEVO) ====================

async def crear_studio_command(update: Update, context: ContextTypes.DEFAULT_TYPE, client=None, config=None):
    """Comando /studio - Activa modo de generaciГіn de estudios."""
    user_id = update.effective_user.id
    # Guardar client y config en context.user_data para usarlos despuГ©s
    context.user_data['client'] = client
    context.user_data['config'] = config
    
    ESPERANDO_PROMPT_STUDIO[user_id] = True
    context.user_data["studio_flow"] = {"step": "tipo", "data": {}}

    opciones_texto = "\n".join(
        f"{indice}. {opcion['label']}" for indice, opcion in enumerate(STUDIO_TIPOS, start=1)
    )
    await update.message.reply_text(
        "рџ“ќ **Modo Studio Activado**\n\n"
        "Primero elige el tipo de informe/estudio (nГєmero o texto):\n"
        f"{opciones_texto}"
    )


async def recibir_prompt_studio(update, context, client, config):
    """Recibe el prompt y lanza el agente de estudio mejorado."""
    user_id = update.effective_user.id
    
    if user_id not in ESPERANDO_PROMPT_STUDIO:
        return False

    studio_flow = context.user_data.get("studio_flow")
    user_text = update.message.text

    if not studio_flow:
        prompt = user_text
    else:
        paso = studio_flow.get("step")
        datos = studio_flow.get("data", {})

        if paso == "tipo":
            seleccion = _seleccionar_opcion(user_text, STUDIO_TIPOS)
            if not seleccion:
                opciones_texto = "\n".join(
                    f"{indice}. {opcion['label']}" for indice, opcion in enumerate(STUDIO_TIPOS, start=1)
                )
                await update.message.reply_text(
                    "No he podido identificar el tipo. Responde con un nГєmero o texto:\n"
                    f"{opciones_texto}"
                )
                return True
            datos["tipo"] = seleccion
            studio_flow["step"] = "tono"
            studio_flow["data"] = datos
            opciones_texto = "\n".join(
                f"{indice}. {opcion['label']}" for indice, opcion in enumerate(STUDIO_TONOS, start=1)
            )
            await update.message.reply_text(
                "Perfecto. Ahora elige el tono del informe:\n"
                f"{opciones_texto}"
            )
            return True

        if paso == "tono":
            seleccion = _seleccionar_opcion(user_text, STUDIO_TONOS)
            if not seleccion:
                opciones_texto = "\n".join(
                    f"{indice}. {opcion['label']}" for indice, opcion in enumerate(STUDIO_TONOS, start=1)
                )
                await update.message.reply_text(
                    "No he podido identificar el tono. Responde con un nГєmero o texto:\n"
                    f"{opciones_texto}"
                )
                return True
            datos["tono"] = seleccion
            studio_flow["step"] = "extension"
            studio_flow["data"] = datos
            opciones_texto = "\n".join(
                f"{indice}. {opcion['label']}" for indice, opcion in enumerate(STUDIO_EXTENSION, start=1)
            )
            await update.message.reply_text(
                "Genial. ВїQuГ© extensiГіn prefieres?\n"
                f"{opciones_texto}"
            )
            return True

        if paso == "extension":
            seleccion = _seleccionar_opcion(user_text, STUDIO_EXTENSION)
            if not seleccion:
                opciones_texto = "\n".join(
                    f"{indice}. {opcion['label']}" for indice, opcion in enumerate(STUDIO_EXTENSION, start=1)
                )
                await update.message.reply_text(
                    "No he podido identificar la extensiГіn. Responde con un nГєmero o texto:\n"
                    f"{opciones_texto}"
                )
                return True
            datos["extension"] = seleccion
            studio_flow["step"] = "nivel"
            studio_flow["data"] = datos
            opciones_texto = "\n".join(
                f"{indice}. {opcion['label']}" for indice, opcion in enumerate(STUDIO_NIVEL, start=1)
            )
            await update.message.reply_text(
                "ВїNivel de profundidad?\n"
                f"{opciones_texto}"
            )
            return True

        if paso == "nivel":
            seleccion = _seleccionar_opcion(user_text, STUDIO_NIVEL)
            if not seleccion:
                opciones_texto = "\n".join(
                    f"{indice}. {opcion['label']}" for indice, opcion in enumerate(STUDIO_NIVEL, start=1)
                )
                await update.message.reply_text(
                    "No he podido identificar el nivel. Responde con un nГєmero o texto:\n"
                    f"{opciones_texto}"
                )
                return True
            datos["nivel"] = seleccion
            studio_flow["step"] = "prompt"
            studio_flow["data"] = datos
            await update.message.reply_text(
                "ВЎListo! Ahora envГ­ame la asignatura y requisitos especГ­ficos.\n"
                "Ejemplo: 'Sistemas InformГЎticos (FP DAW), prГЎctica sobre redes con enunciado, guГ­a paso a paso y resoluciГіn completa'."
            )
            return True

        if paso == "prompt":
            prompt = user_text
        else:
            prompt = user_text

    datos = (studio_flow or {}).get("data", {})
    preferencias = [
        f"- Tipo: {datos.get('tipo', {}).get('label', 'No especificado')}",
        f"- Tono: {datos.get('tono', {}).get('label', 'No especificado')}",
        f"- ExtensiГіn: {datos.get('extension', {}).get('label', 'No especificado')}",
        f"- Profundidad: {datos.get('nivel', {}).get('label', 'No especificado')}",
    ]
    prompt_final = (
        f"{prompt}\n\nPreferencias:\n" + "\n".join(preferencias)
    )

    context.user_data.pop("studio_flow", None)
    del ESPERANDO_PROMPT_STUDIO[user_id]

    await update.message.reply_text(
        "рџљЂ **Agente Studio V2 iniciado**\n\n"
        "Voy a generar un documento profesional con:\n"
        "вњ… Estructura optimizada\n"
        "вњ… Contenido denso y especГ­fico\n"
        "вњ… ValidaciГіn de calidad\n\n"
        "Te irГ© informando del progreso. Esto puede tardar varios minutos..."
    )

    # CRГЌTICO: context.application.create_task para que se ejecute en background
    context.application.create_task(
        agente_estudio_mejorado(
            prompt_final,
            update.effective_chat.id,
            context,
            config,
            client,
            preferencias=datos,
        )
    )

    return True



# async def agente_estudio_mejorado(prompt_usuario, chat_id, context, config, client):
#     """
#     Agente de estudio mejorado usando ContentEngine con reporting de progreso.
#     """
#     try:
#         # DETECCIГ“N DE TESTS SIMPLES
#         palabras_test = ['test', 'prueba', '2+2', 'hola mundo', 'ejemplo', 'demo']
#         es_test = any(palabra in prompt_usuario.lower() for palabra in palabras_test)
        
#         if es_test and len(prompt_usuario.split()) < 10:
#             await context.bot.send_message(
#                 chat_id=chat_id,
#                 text="рџ¤– **Detecto que esto es una prueba/test.**\n\n"
#                      "Para un estudio acadГ©mico completo, usa temas reales como:\n"
#                      "вЂў 'ProgramaciГіn didГЎctica de Bases de Datos en DAW'\n"
#                      "вЂў 'TFG sobre ciberseguridad en banca'\n"
#                      "вЂў 'Estudio sobre IA en medicina'\n\n"
#                      "Si quieres continuar con este test de todos modos, "
#                      "vuelve a enviarlo precedido de 'forzar:'\n"
#                      "Ejemplo: `forzar: test 2+2`"
#             )
#             return
        
#         # Limpiar prefijo 'forzar:' si existe
#         if prompt_usuario.lower().startswith('forzar:'):
#             prompt_usuario = prompt_usuario[7:].strip()
        
#         # Parsear prompt para detectar tipo
#         tipo = "general"
#         if "programaciГіn didГЎctica" in prompt_usuario.lower() or "programacion didactica" in prompt_usuario.lower():
#             tipo = "programacion_didactica"
#         elif "tfg" in prompt_usuario.lower() or "tfm" in prompt_usuario.lower():
#             tipo = "tfg"
#         elif "investigaciГіn" in prompt_usuario.lower() or "investigacion" in prompt_usuario.lower():
#             tipo = "investigacion"
        
#         # FASE 1: Generar estructura
#         await context.bot.send_message(chat_id=chat_id, text="рџ§  **Fase 1: Analizando tema y generando estructura...**")
        
#         from content_engine import ContentEngine
#         engine = ContentEngine(client, modelo_avanzado=MODELO_GENERACION)

#         # Generar solo la estructura primero
#         # Si el prompt es muy largo y especГ­fico, debemos forzar al engine a no usar su template estГЎndar. 
#         if len(prompt_usuario.split()) > 150: # Contamos palabras, no caracteres
#             prompt_bypass = f"""
#             Genera un Г­ndice para este proyecto: {prompt_usuario}
#             REGLAS:
#             - Responde SOLO con los tГ­tulos de las secciones.
#             - MГЎximo 10 secciones.
#             - Una secciГіn por lГ­nea.
#             - No escribas introducciones ni despedidas.
#             """
#             res = client.chat.complete(
#                 model=MODELO_GENERACION,
#                 messages=[{"role": "user", "content": prompt_bypass}]
#             )
#             contenido_respuesta = res.choices[0].message.content
            
#             # Limpieza robusta:
#             # 1. Dividimos por lГ­neas
#             lineas = contenido_respuesta.split('\n')
#             secciones = []
#             for l in lineas:
#                 # Quitamos nГєmeros, guiones y puntos al principio (ej: "1. TГ­tulo" -> "TГ­tulo")
#                 limpia = re.sub(r'^[\d\.\-\s]+', '', l).strip()
#                 if limpia and len(limpia) > 3: # Evitamos lГ­neas vacГ­as o basura
#                     secciones.append(limpia)
            
#             estructura = {'secciones': secciones}
#         else:
#             estructura = await engine._generar_estructura(prompt_usuario, tipo, "universitario", "completo")
        
#         await context.bot.send_message(
#             chat_id=chat_id, 
#             text=f"рџ“‹ Estructura generada: {len(estructura['secciones'])} secciones\n\n"
#                  f"**Fase 2: Redactando contenido profesional...**\n"
#                  f"(Esto llevarГЎ varios minutos, te voy informando)"
#         )
        
#         # Crear documento Word
#         timestamp = datetime.now().strftime('%Y%m%d_%H%M')
#         safe_prompt = re.sub(r'[^a-zA-Z0-9]', '', prompt_usuario[:20])
#         nombre_carpeta = f"STUDIO_{timestamp}_{safe_prompt}"
#         ruta_base = f"/app/documentos/EstudiosRobi/{nombre_carpeta}"
#         os.makedirs(ruta_base, exist_ok=True)
#         ruta_archivo = os.path.join(ruta_base, "estudio_completo.docx")
        
#         # FASE 2: Desarrollar cada secciГіn CON PROGRESO
#         secciones_desarrolladas = []
        
#         for i, seccion in enumerate(estructura['secciones'], 1):
#             await context.bot.send_message(chat_id=chat_id, text=f"вњЌпёЏ **SecciГіn {i}/{len(estructura['secciones'])}:** {seccion}")
            
#             # 1. Generar contenido (crudo con JSON)
#             contenido_raw = await engine._desarrollar_seccion(
#                 seccion=seccion,
#                 tema_global=prompt_usuario,
#                 contexto_previo=secciones_desarrolladas,
#                 numero=i,
#                 total=len(estructura['secciones'])
#             )
            
#             # 1. EXTRAER EL GRГЃFICO (IMPORTANTE: Antes de limpiar)
#             # Buscamos si en ESTA secciГіn la IA ha metido un JSON
#             datos_v = engine._extraer_datos_visuales(contenido_raw)
            
#             # 2. LIMPIAR EL TEXTO
#             # Eliminamos las etiquetas [GRAFICO_DATA] para que no salgan escritas en el Word
#             contenido_final = engine._limpiar_formato(contenido_raw)

#             # 3. REFINAR (opcional)
#             if await engine._necesita_refinamiento(contenido_final):
#                 contenido_final = await engine._refinar_contenido(contenido_final, seccion)

#             # 4. GUARDAR EN LA LISTA CON SUS DATOS VISUALES
#             secciones_desarrolladas.append({
#                 'titulo': seccion,
#                 'contenido': contenido_final,
#                 'numero': i,
#                 'datos_visuales': datos_v  # <-- Si hay JSON, aquГ­ se guarda
#             })
            
#             await context.bot.send_message(chat_id=chat_id, text=f"вњ… SecciГіn {i} lista.")
#             if i < len(estructura['secciones']): await asyncio.sleep(8)

#         # --- AQUГЌ ESTГЃ LA MAGIA ---
        
#         # 6. Preparar el objeto de datos para el exportador
#         estudio_data = {
#             'indice': [s['titulo'] for s in secciones_desarrolladas],
#             'secciones': secciones_desarrolladas,
#             'metadata': {
#                 'tema': prompt_usuario,
#                 'nivel': "universitario"
#             }
#         }

#         # 7. Generar el Word Real con GrГЎficos
#         exportar_a_word_premium(estudio_data, ruta_archivo)
        
#         # 8. Enviar el archivo final
#         total_palabras = sum(len(s['contenido'].split()) for s in secciones_desarrolladas)
#         with open(ruta_archivo, "rb") as doc_file:
#             await context.bot.send_document(
#                 chat_id=chat_id,
#                 document=doc_file,
#                 caption=f"рџЋ‰ **ВЎEstudio completado!**\n\nрџ“Љ Secciones: {len(secciones_desarrolladas)}\nрџ“Џ ~{total_palabras:,} palabras\nрџ“€ GrГЎficos e imГЎgenes incluidos."
#             )

#     except Exception as e:
#         logging.error(f"Error en agente_estudio_mejorado: {e}")
#         await context.bot.send_message(chat_id=chat_id, text=f"вќЊ Error: {str(e)}")

def _resolver_tipo_estudio(prompt_usuario: str, preferencias: dict | None) -> str:
    tipo_id = (preferencias or {}).get("tipo", {}).get("id")
    prompt_lower = prompt_usuario.lower()

    if tipo_id == "practica_asignatura":
        return "practica_asignatura"

    if tipo_id == "investigacion_mercado":
        return "investigacion"

    if tipo_id == "estudio_academico":
        if "programaciГіn didГЎctica" in prompt_lower or "programacion didactica" in prompt_lower:
            return "programacion_didactica"
        if "tfg" in prompt_lower or "tfm" in prompt_lower:
            return "tfg"
        if "investigaciГіn" in prompt_lower or "investigacion" in prompt_lower:
            return "investigacion"

    return "general"


def _resolver_extension_estudio(preferencias: dict | None) -> str:
    extension_id = (preferencias or {}).get("extension", {}).get("id")
    return {
        "corto": "breve",
        "medio": "medio",
        "largo": "completo",
        "extenso": "extenso",
    }.get(extension_id, "breve")


def _resolver_nivel_estudio(preferencias: dict | None) -> str:
    nivel_label = (preferencias or {}).get("nivel", {}).get("label")
    return nivel_label or "universitario"


def _resolver_tono_estudio(preferencias: dict | None) -> str:
    tono_label = (preferencias or {}).get("tono", {}).get("label")
    return tono_label or "formal"


def _resolver_perfil_redactor(preferencias: dict | None) -> str:
    tipo_id = (preferencias or {}).get("tipo", {}).get("id")
    if tipo_id == "practica_asignatura":
        return "docente de FP InformГЎtica"
    tipo_label = (preferencias or {}).get("tipo", {}).get("label")
    return tipo_label or "acadГ©mico"


async def agente_estudio_mejorado(prompt_usuario, chat_id, context, config, client, preferencias=None):
    """
    Agente de estudio mejorado con generaciГіn REAL de grГЎficos.
    """
    try:
        # DETECCIГ“N DE TESTS SIMPLES
        palabras_test = ['test', 'prueba', '2+2', 'hola mundo', 'ejemplo', 'demo']
        es_test = any(palabra in prompt_usuario.lower() for palabra in palabras_test)
        
        if es_test and len(prompt_usuario.split()) < 10:
            await context.bot.send_message(
                chat_id=chat_id,
                text="рџ¤– **Detecto que esto es una prueba/test.**\n\n"
                     "Para un estudio acadГ©mico completo, usa temas reales como:\n"
                     "вЂў 'ProgramaciГіn didГЎctica de Bases de Datos en DAW'\n"
                     "вЂў 'TFG sobre ciberseguridad en banca'\n"
                     "вЂў 'Estudio sobre IA en medicina'\n\n"
                     "Si quieres continuar con este test de todos modos, "
                     "vuelve a enviarlo precedido de 'forzar:'\n"
                     "Ejemplo: `forzar: test 2+2`"
            )
            return
        
        # Limpiar prefijo 'forzar:' si existe
        if prompt_usuario.lower().startswith('forzar:'):
            prompt_usuario = prompt_usuario[7:].strip()
        
        tipo = _resolver_tipo_estudio(prompt_usuario, preferencias)
        extension = _resolver_extension_estudio(preferencias)
        nivel = _resolver_nivel_estudio(preferencias)
        tono = _resolver_tono_estudio(preferencias)
        perfil_redactor = _resolver_perfil_redactor(preferencias)
        
        # FASE 1: Generar estructura
        await context.bot.send_message(chat_id=chat_id, text="рџ§  **Fase 1: Analizando tema y generando estructura...**")
        
        from content_engine import ContentEngine
        engine = ContentEngine(
            client,
            modelo_avanzado=MODELO_GENERACION,
            perfil_redactor=perfil_redactor,
            tono=tono,
        )
        
        estructura = await engine._generar_estructura(prompt_usuario, tipo, nivel, extension)
        
        await context.bot.send_message(
            chat_id=chat_id, 
            text=f"рџ“‹ Estructura generada: {len(estructura['secciones'])} secciones\n\n"
                 f"**Fase 2: Redactando contenido con anГЎlisis de datos para grГЎficos...**\n"
                 f"(Esto llevarГЎ varios minutos)"
        )
        
        # Crear documento Word
        timestamp = datetime.now().strftime('%Y%m%d_%H%M')
        safe_prompt = re.sub(r'[^a-zA-Z0-9]', '', prompt_usuario[:20])
        nombre_carpeta = f"STUDIO_{timestamp}_{safe_prompt}"
        ruta_base = f"/app/documentos/EstudiosRobi/{nombre_carpeta}"
        os.makedirs(ruta_base, exist_ok=True)
        ruta_archivo = os.path.join(ruta_base, "estudio_completo.docx")
        
        from docx import Document
        from docx.shared import Pt, Inches
        doc = Document()
        
        # Portada
        doc.add_heading(f'ESTUDIO: {prompt_usuario[:80]}', 0)
        doc.add_paragraph(f"Generado el: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
        tipo_label = (preferencias or {}).get("tipo", {}).get("label")
        doc.add_paragraph(f"Tipo: {tipo_label or tipo.replace('_', ' ').title()}")
        doc.add_paragraph(f"Tono: {tono}")
        doc.add_paragraph(f"Nivel: {nivel}")
        doc.add_paragraph(f"ExtensiГіn: {extension}")
        doc.add_paragraph(f"Modelo: {MODELO_GENERACION}")
        doc.add_page_break()
        
        # ГЌndice
        doc.add_heading('ГЌNDICE', 1)
        for i, titulo in enumerate(estructura['secciones'], 1):
            doc.add_paragraph(f"{i}. {titulo}", style='List Number')
        doc.add_page_break()
        
        # FASE 2: Desarrollar secciones
        secciones_desarrolladas = []
        
        for i, seccion in enumerate(estructura['secciones'], 1):
            await context.bot.send_message(
                chat_id=chat_id, 
                text=f"вњЌпёЏ **SecciГіn {i}/{len(estructura['secciones'])}:** {seccion}\n\n_Generando contenido..._"
            )
            
            # Generar contenido
            max_intentos = 3
            contenido = None
            
            for intento in range(max_intentos):
                try:
                    contenido = await engine._desarrollar_seccion(
                        seccion=seccion,
                        tema_global=prompt_usuario,
                        contexto_previo=secciones_desarrolladas,
                        numero=i,
                        total=len(estructura['secciones']),
                        extension=extension
                    )
                    break
                except Exception as e:
                    if "429" in str(e) and intento < max_intentos - 1:
                        espera = 10 * (intento + 1)
                        await context.bot.send_message(
                            chat_id=chat_id,
                            text=f"вЏі Rate limit. Reintentando en {espera}s..."
                        )
                        await asyncio.sleep(espera)
                    elif intento == max_intentos - 1:
                        contenido = "[Contenido no disponible por error de API]"
            
            # Validar y refinar
            if contenido and await engine._necesita_refinamiento(contenido, extension):
                await context.bot.send_message(
                    chat_id=chat_id,
                    text=f"рџ”„ Refinando contenido..."
                )
                try:
                    contenido = await engine._refinar_contenido(
                        contenido,
                        seccion,
                        extension=extension
                    )
                except:
                    pass
            
            # Guardar secciГіn
            secciones_desarrolladas.append({
                'titulo': seccion,
                'contenido': contenido,
                'numero': i
            })
            
            await context.bot.send_message(
                chat_id=chat_id,
                text=f"вњ… SecciГіn {i} completada ({len(contenido.split())} palabras)"
            )
            
            # Pausa anti-rate-limit
            if i < len(estructura['secciones']):
                await asyncio.sleep(8)
        
        # FASE 3: GENERAR DOCUMENTO CON GRГЃFICOS
        await context.bot.send_message(
            chat_id=chat_id,
            text="рџ“Љ **Fase 3: Analizando secciones para generar grГЎficos...**\n\n"
                 "Esto puede tardar unos minutos adicionales."
        )
        
        # Importar el generador de grГЎficos
        from generador_graficos import IntegradorGraficosWord
        
        integrador = IntegradorGraficosWord(client, modelo=MODELO_GENERACION)
        total_visuales = 0
        
        # Procesar cada secciГіn con el integrador de grГЎficos
        for seccion in secciones_desarrolladas:
            try:
                await context.bot.send_message(
                    chat_id=chat_id,
                    text=f"рџ”Ќ Analizando datos en: {seccion['titulo'][:40]}..."
                )
                
                resultado = await integrador.procesar_seccion_con_graficos(
                    titulo_seccion=seccion['titulo'],
                    contenido=seccion['contenido'],
                    doc=doc,
                    numero_seccion=seccion['numero']
                )
                
                total_visuales += resultado["total"]
                
                if resultado["total"] > 0:
                    await context.bot.send_message(
                        chat_id=chat_id,
                        text=f"вњ… {resultado['total']} recurso(s) visual(es) aГ±adido(s) en secciГіn {seccion['numero']}"
                    )
            
            except Exception as e:
                logging.error(f"Error generando grГЎficos en secciГіn {seccion['numero']}: {e}")
                # Continuar con las demГЎs secciones
                continue
        
        # Si no se aГ±adieron grГЎficos automГЎticamente, aГ±adir las secciones como texto
        if total_visuales == 0:
            await context.bot.send_message(
                chat_id=chat_id,
                text="в„№пёЏ No se detectaron datos o modelos visualizables. Generando documento estructurado."
            )
            
            # Recorrer secciones y aГ±adir al documento con mejor estructura
            for seccion in secciones_desarrolladas:
                doc.add_heading(seccion['titulo'], level=1)
                
                # Estructurar el contenido en puntos y subpuntos
                contenido_estructurado = _estructurar_contenido_en_puntos(seccion['contenido'])
                _agregar_contenido_estructurado_al_doc(doc, contenido_estructurado)
                
                doc.add_page_break()
        
        # Guardar documento
        doc.save(ruta_archivo)
        
        # Enviar resultado
        total_palabras = sum(len(s['contenido'].split()) for s in secciones_desarrolladas)
        
        mensaje_final = f"рџЋ‰ **ВЎEstudio completado!**\n\n"
        mensaje_final += f"рџ“„ Archivo: `{ruta_archivo}`\n"
        mensaje_final += f"рџ“Љ Secciones: {len(secciones_desarrolladas)}\n"
        mensaje_final += f"рџ“Џ ExtensiГіn: ~{total_palabras:,} palabras\n"
        
        if total_visuales > 0:
            mensaje_final += f"рџ“€ Recursos visuales generados: {total_visuales}\n"
        
        mensaje_final += f"\nEl documento estГЎ listo en tu carpeta de Documentos."
        
        await context.bot.send_message(chat_id=chat_id, text=mensaje_final)
        
        # Opcional: Enviar el archivo directamente por Telegram
        try:
            with open(ruta_archivo, 'rb') as doc_file:
                await context.bot.send_document(
                    chat_id=chat_id,
                    document=doc_file,
                    caption=f"рџ“„ Estudio: {prompt_usuario[:50]}..."
                )
        except Exception as e:
            logging.error(f"Error enviando documento: {e}")
    
    except Exception as e:
        logging.error(f"Error en agente_estudio_mejorado: {e}")
        import traceback
        traceback.print_exc()
        await context.bot.send_message(
            chat_id=chat_id,
            text=f"вќЊ **Error generando estudio**\n\n{str(e)}\n\n"
                 f"Prueba con un tema mГЎs especГ­fico o contacta con el administrador."
        )


# ==================== FUNCIONES AUXILIARES ====================

def _estructurar_contenido_en_puntos(contenido: str) -> list:
    """
    Convierte un bloque de texto en una estructura de puntos y subpuntos.
    
    Returns:
        Lista de diccionarios con 'tipo' ('titulo', 'punto', 'subpunto', 'parrafo') y 'texto'
    """
    elementos = []
    lineas = contenido.split('\n')

    def _fragmentar_texto_largo(texto: str) -> list[str]:
        """Divide texto largo en bloques cortos para evitar pГЎrrafos de varias pГЎginas."""
        frases = re.split(r'(?<=[.!?])\s+', texto)
        bloques = []
        bloque_actual = ""

        for frase in frases:
            frase = frase.strip()
            if not frase:
                continue

            candidato = f"{bloque_actual} {frase}".strip()
            if len(candidato) <= 240:
                bloque_actual = candidato
                continue

            if bloque_actual:
                bloques.append(bloque_actual)
            bloque_actual = frase

        if bloque_actual:
            bloques.append(bloque_actual)

        return bloques or [texto]
    
    for linea in lineas:
        linea_limpia = linea.strip()
        
        if not linea_limpia:
            continue
        
        # Detectar tГ­tulos (lГ­neas cortas que terminan en : o son todo mayГєsculas)
        if (len(linea_limpia) < 80 and linea_limpia.endswith(':')) or \
           (len(linea_limpia) < 80 and linea_limpia.isupper() and len(linea_limpia.split()) <= 6):
            elementos.append({
                'tipo': 'titulo',
                'texto': linea_limpia.rstrip(':')
            })
            continue
        
        # Detectar subpuntos (sangrГ­a + marcador) usando la lГ­nea original
        if re.match(r'^\s{2,}[-вЂўв–Єo]\s+', linea) or \
           re.match(r'^\s{2,}[\d]+[\.\)]\s+', linea):
            texto_limpio = re.sub(r'^\s+[-вЂўв–Єo]\s+|^\s+[\d]+[\.\)]\s+', '', linea).strip()
            elementos.append({
                'tipo': 'subpunto',
                'texto': texto_limpio
            })
            continue

        # Detectar puntos que empiezan con nГєmeros, guiones o marcadores
        if re.match(r'^[\d]+[\.\)]\s+', linea_limpia) or \
           re.match(r'^[-вЂўв–Є]\s+', linea_limpia):
            # Es un punto principal
            texto_limpio = re.sub(r'^[\d]+[\.\)]\s+|^[-вЂўв–Є]\s+', '', linea_limpia)
            elementos.append({
                'tipo': 'punto',
                'texto': texto_limpio
            })
            continue
        
        # Si la lГ­nea es muy corta (< 100 caracteres) y termina en punto, puede ser un punto
        if len(linea_limpia) < 100 and linea_limpia.endswith('.') and not linea_limpia.endswith('...'):
            # Contar si hay muchas mayГєsculas al principio (indicador de tГ­tulo/punto)
            palabras = linea_limpia.split()
            if palabras and palabras[0][0].isupper():
                elementos.append({
                    'tipo': 'punto',
                    'texto': linea_limpia
                })
                continue
        
        # Todo lo demГЎs es pГЎrrafo normal, con corte automГЎtico si es largo
        if len(linea_limpia) > 220:
            for bloque in _fragmentar_texto_largo(linea_limpia):
                elementos.append({
                    'tipo': 'punto' if len(bloque) < 170 else 'parrafo',
                    'texto': bloque
                })
        else:
            elementos.append({
                'tipo': 'parrafo',
                'texto': linea_limpia
            })
    
    return elementos


def _agregar_contenido_estructurado_al_doc(doc, elementos: list):
    """
    AГ±ade elementos estructurados al documento de Word.
    
    Args:
        doc: Documento de python-docx
        elementos: Lista de diccionarios con tipo y texto
    """
    for elemento in elementos:
        tipo = elemento['tipo']
        texto = elemento['texto']
        
        if tipo == 'titulo':
            # AГ±adir como subtГ­tulo (heading nivel 2)
            doc.add_heading(texto, level=2)
        
        elif tipo == 'punto':
            # AГ±adir como viГ±eta de nivel 1
            p = doc.add_paragraph(texto, style='List Bullet')
        
        elif tipo == 'subpunto':
            # AГ±adir como viГ±eta de nivel 2
            p = doc.add_paragraph(texto, style='List Bullet 2')
        
        else:  # parrafo
            # PГЎrrafo normal, pero si es corto, convertir a punto
            if len(texto) < 200:
                doc.add_paragraph(texto, style='List Bullet')
            else:
                doc.add_paragraph(texto)


async def ejecutar_evaluacion_cartera(update, client, config):
    """
    Ejecuta la evaluaciГіn completa de cartera y envГ­a el resultado.
    Se ejecuta en background para no bloquear el bot.
    """
    try:
        ruta_excel = "/app/documentos/bolsav2.xlsx"
        
        # Validar que existe el archivo
        import os
        if not os.path.exists(ruta_excel):
            await update.message.reply_text(
                "вќЊ No encuentro tu archivo de inversiones.\n\n"
                "Comprueba que:\n"
                "вЂў El volumen estГЎ montado\n"
                "вЂў El archivo se llama 'bolsav2.xlsx'\n"
                "вЂў EstГЎ en la ruta correcta"
            )
            return
        
        # Crear evaluador
        from tools_system import buscar_internet
        evaluador = EvaluadorProfesionalCartera(
            client=client,
            buscar_internet=lambda q: buscar_internet(q, client, config, MODELO_GENERACION),
            modelo=MODELO_GENERACION
        )
        
        # Evaluar cartera
        logging.info("рџ”Ќ Iniciando evaluaciГіn de cartera...")
        resultado = await evaluador.evaluar_cartera_completa(ruta_excel)
        
        # Formatear resultado
        mensaje = formatear_informe_profesional(resultado)
        
        # Enviar resultado (dividir si es muy largo)
        await enviar_mensaje_largo(update, mensaje)
        
        logging.info("вњ… EvaluaciГіn de cartera completada")
    
    except Exception as e:
        logging.error(f"Error en evaluaciГіn de cartera: {e}")
        import traceback
        traceback.print_exc()
        await update.message.reply_text(
            f"вќЊ Error al evaluar cartera: {str(e)}\n\n"
            "IntГ©ntalo de nuevo en unos minutos."
        )



async def ejecutar_evaluacion_valor(update, valor, client, config):
    """Ejecuta evaluaciГіn profesional de un valor y envГ­a resultado."""
    try:
        from tools_system import buscar_internet
        evaluador = EvaluadorProfesionalCartera(
            client=client,
            buscar_internet=lambda q: buscar_internet(q, client, config, MODELO_GENERACION),
            modelo=MODELO_GENERACION
        )
        resultado = await evaluador.evaluar_valor_unico(valor)
        mensaje = formatear_informe_profesional(resultado)
        await enviar_mensaje_largo(update, mensaje)
    except Exception as e:
        logging.error(f"Error en evaluaciГіn de valor: {e}")
        await update.message.reply_text(f"вќЊ Error: {str(e)}")


async def lanzar_escaner_oportunidades(update, mercado, client, config):
    """Ejecuta bГєsqueda de oportunidades de inversiГіn con anГЎlisis profesional."""
    try:
        await update.message.reply_text(
            f"рџ”Ќ Escaneando {mercado.upper()} con anГЎlisis multi-capa...\n\n"
            "Esto llevarГЎ varios minutos mientras evalГєo cada candidato. Te aviso cuando termine."
        )
        
        # Crear evaluador profesional
        from tools_system import buscar_internet
        from evaluador_profesional import EvaluadorProfesionalCartera
        
        evaluador = EvaluadorProfesionalCartera(
            client=client,
            buscar_internet=lambda q: buscar_internet(q, client, config, MODELO_GENERACION),
            modelo=MODELO_GENERACION
        )
        
        # Ejecutar bГєsqueda con evaluador
        resultado = await buscar_oportunidades_inversion(
            mercado=mercado,
            client=client,
            buscar_internet=lambda q: buscar_internet(q, client, config, MODELO_GENERACION),
            evaluador=evaluador  # <- NUEVO: Pasa el evaluador para anГЎlisis completo
        )
        
        await enviar_mensaje_largo(update, resultado)
        
    except Exception as e:
        logging.error(f"Error en oportunidades: {e}")
        import traceback
        traceback.print_exc()
        await update.message.reply_text(f"вќЊ Error al escanear oportunidades: {str(e)}")